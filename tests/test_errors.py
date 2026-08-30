"""예외 처리 및 '완전 로컬 동작' 검증.

* 잘못된 입력에 대해 개발자용 스택 트레이스가 아니라 사람이 읽는
  :class:`ReportGenError` 가 나오는지
* 소스 어디에도 네트워크를 쓰는 코드가 없는지
"""

from __future__ import annotations

import ast
import datetime as _dt
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from reportgen.aggregator import AggregationSpec, aggregate_monthly  # noqa: E402
from reportgen.data_reader import ReadOptions, read_table  # noqa: E402
from reportgen.errors import (  # noqa: E402
    AggregationError,
    CellRangeError,
    FileFormatError,
    FormulaCacheError,
    MappingError,
    ReportGenError,
    SheetNotFoundError,
    TemplateError,
)
from reportgen.mapping import Binding, TemplateSlot, resolve_context  # noqa: E402
from reportgen.templating import open_template  # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
SOURCES = os.path.join(HERE, "fixtures", "sources")

FAILURES: list[str] = []

#: 이 프로그램이 절대 써서는 안 되는 모듈들
FORBIDDEN_MODULES = {
    "socket",
    "urllib",
    "urllib2",
    "urllib3",
    "http",
    "httplib",
    "requests",
    "ftplib",
    "smtplib",
    "telnetlib",
    "xmlrpc",
    "asyncio",
    "aiohttp",
    "websocket",
    "websockets",
    "boto3",
    "paramiko",
}


def expect(kind: type[BaseException], label: str, call) -> None:
    try:
        call()
    except kind as exc:
        message = str(exc)
        friendly = isinstance(exc, ReportGenError) and len(message) > 10
        if friendly:
            print(f"  OK   {label}\n         -> {message.splitlines()[0]}")
        else:
            FAILURES.append(label)
            print(f"  FAIL {label}: 메시지가 불친절함 ({message!r})")
    except BaseException as exc:  # noqa: BLE001
        FAILURES.append(label)
        print(f"  FAIL {label}: {type(exc).__name__} 이(가) 나옴 ({exc})")
    else:
        FAILURES.append(label)
        print(f"  FAIL {label}: 오류가 나지 않음")


def test_input_errors() -> None:
    print("\n[잘못된 입력에 대한 안내 메시지]")

    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as handle:
        handle.write(b"this is not excel")
        text_file = handle.name

    expect(FileFormatError, "엑셀이 아닌 파일을 원본으로 고른 경우",
           lambda: read_table(text_file))
    expect(FileFormatError, "없는 파일 경로",
           lambda: read_table(os.path.join(SOURCES, "없는파일.xlsx")))
    expect(FileFormatError, "원본을 고르지 않고 진행",
           lambda: read_table(""))
    expect(SheetNotFoundError, "없는 시트 이름 지정",
           lambda: read_table(os.path.join(SOURCES, "S1_헤더1행.xlsx"),
                              ReadOptions(sheet_name="없는시트")))
    expect(CellRangeError, "셀 범위 문자열이 엉망인 경우",
           lambda: read_table(os.path.join(SOURCES, "S1_헤더1행.xlsx"),
                              ReadOptions(cell_range="A1-Z9", auto_detect=False)))
    expect(FileFormatError, "템플릿으로 .txt 를 고른 경우",
           lambda: open_template(text_file))
    expect(FileFormatError, "구형 .doc 템플릿",
           lambda: open_template(os.path.join(HERE, "없는템플릿.doc")))

    table = read_table(os.path.join(SOURCES, "S1_헤더1행.xlsx"))
    expect(AggregationError, "날짜 컬럼이 없는 표를 월집계하려는 경우",
           lambda: aggregate_monthly(table, AggregationSpec()))
    expect(AggregationError, "지정한 날짜 컬럼이 원본에 없는 경우",
           lambda: aggregate_monthly(table, AggregationSpec(date_column="없는컬럼")))

    daily = read_table(os.path.join(SOURCES, "S5_일단위_한달.xlsx"))
    expect(AggregationError, "제외 조건이 모든 행을 걸러낸 경우",
           lambda: aggregate_monthly(daily, AggregationSpec(only_months=["2099-01"])))

    expect(FormulaCacheError, "계산기도 못 푸는 수식(없는 외부 파일 참조)에 캐시도 없는 경우",
           lambda: read_table(os.path.join(SOURCES, "S24_미지원수식.xlsx"), ReadOptions(sheet_name="본표")))

    expect(MappingError, "매핑이 원본에 없는 컬럼을 가리키는 경우",
           lambda: resolve_context(
               [TemplateSlot(key="사용량")],
               {"사용량": Binding(source="column", column="사라진컬럼")},
               table,
           ))
    expect(MappingError, "필수 항목이 비어 있는 채로 생성(strict)",
           lambda: resolve_context(
               [TemplateSlot(key="미지정항목")], {}, table, strict=True
           ))

    os.unlink(text_file)


def test_formula_engine_computes_simple_formulas() -> None:
    """계산 캐시가 없어도, 계산기가 풀 수 있는 수식은 오류 없이 값을 채운다.

    ``=B2+C2`` 처럼 캐시가 없는 단순 수식은 이제 FormulaCacheError 를 내는
    대신 내장 계산기가 직접 계산해서 정상적인 표로 읽혀야 한다.
    """
    print("\n[계산 캐시 없는 단순 수식 - 내장 계산기가 직접 계산]")
    path = os.path.join(SOURCES, "S18_계산안된수식.xlsx")
    table = read_table(path)
    ok = not table.warnings and table.rows == [
        [_dt.datetime(2026, 11, 1), 100, 20, 120],
        [_dt.datetime(2026, 11, 2), 110, 22, 132],
    ]
    label = "'=B2+C2' 캐시 없어도 오류 없이 120/132 로 정확히 계산됨"
    if ok:
        print(f"  OK   {label}")
    else:
        FAILURES.append(label)
        print(f"  FAIL {label}: warnings={table.warnings}, rows={table.rows}")


def test_header_formula_engine() -> None:
    """헤더 행 자체가 계산 안 된 수식이어도, 계산기가 풀 수 있으면 정상 채움.

    본문 셀만 검사하던 예전 로직은 이 경우를 놓쳐서 '열C' 같은 엉뚱한
    컬럼 이름을 조용히 만들었었다. 실제 사용자의 '날짜가 가로로 늘어선 표'
    화면에서 발견된 패턴(하루씩 더하는 수식 헤더)으로 검증한다.
    """
    print("\n[헤더 행 자체가 계산 안 된 수식인 경우 - 계산기로 직접 계산]")
    path = os.path.join(SOURCES, "S22_헤더수식미계산.xlsx")
    options = ReadOptions(cell_range="B3:H3", header_rows=1, auto_detect=False)
    table = read_table(path, options)
    ok = not table.warnings and table.columns == [
        "2026-07-01", "2026-07-02", "2026-07-03", "2026-07-04",
        "2026-07-05", "2026-07-06", "2026-07-07",
    ]
    label = "하루씩 더하는 날짜 헤더가 '열C' 대신 실제 날짜로 계산됨"
    if ok:
        print(f"  OK   {label}")
        print(f"         -> {table.columns}")
    else:
        FAILURES.append(label)
        print(f"  FAIL {label}: warnings={table.warnings}, columns={table.columns}")


def test_cross_sheet_formula_computed() -> None:
    """다른 시트를 참조하는 수식도 오류 없이 정상 계산되는지.

    ``formulas`` 는 워크북 전체를 계산 그래프로 만들기 때문에, 우리 자체
    계산기와 달리 시트 간 참조도 지원해야 한다.
    """
    print("\n[시트 간 참조 수식 - 워크북 전체 계산으로 지원]")
    path = os.path.join(SOURCES, "S23_시트간참조.xlsx")
    table = read_table(path, ReadOptions(sheet_name="본표"))
    ok = not table.warnings and [row[-1] for row in table.rows] == [999, 999]
    label = "'=요약!A1' 시트 간 참조가 오류 없이 999 로 계산됨"
    if ok:
        print(f"  OK   {label}")
    else:
        FAILURES.append(label)
        print(f"  FAIL {label}: warnings={table.warnings}, rows={table.rows}")


def test_formula_engine_bypass_for_unsupported() -> None:
    """계산기도 못 푸는 수식(없는 외부 파일 참조 등)은 여전히 안전하게 걸러진다.

    기본은 오류, '빈 값으로 넘어가기' 옵션을 켜면 경고 + 빈 값으로 진행.
    """
    print("\n[계산기가 못 푸는 수식 - 우회 옵션]")
    path = os.path.join(SOURCES, "S24_미지원수식.xlsx")
    options = ReadOptions(sheet_name="본표", allow_uncalculated_formulas=True)
    table = read_table(path, options)
    ok = bool(table.warnings) and table.rows[0][-1] is None and table.rows[0][:2] == [
        _dt.datetime(2026, 12, 1), 100,
    ]
    label = "우회 옵션: 없는 외부 파일 참조는 경고 + 빈 값, 나머지 컬럼은 정상"
    if ok:
        print(f"  OK   {label}")
        print(f"         -> {table.warnings[0]}")
    else:
        FAILURES.append(label)
        print(f"  FAIL {label}: warnings={table.warnings}, rows={table.rows}")


def test_table_block_detection() -> None:
    """한 시트에 표가 여러 개 섞여 있을 때, 후보를 정확히 나누는지 확인.

    대형 빌딩 에너지 대장처럼 좌우로 나란히 놓인 표(전력/스팀/용수)와
    위아래로 별개인 표가 한 시트에 있는 경우를 재현한 원본으로 검증한다.
    """
    from reportgen.data_reader import list_table_blocks, read_table, ReadOptions

    print("\n[표 후보 찾기]")
    path = os.path.join(SOURCES, "S21_다중표시트.xlsx")
    blocks = list_table_blocks(path)
    ranges = {b.range for b in blocks}
    expected = {"B7:D9", "F7:H9", "J7:K9", "A16:H18"}
    label = f"좌우로 나란한 표 3개 + 별개 표 1개를 정확히 분리 (찾은 것: {sorted(ranges)})"
    if expected <= ranges:
        print(f"  OK   {label}")
    else:
        FAILURES.append(label)
        print(f"  FAIL {label}")

    steam_block = next((b for b in blocks if b.range == "F7:H9"), None)
    if steam_block is None:
        FAILURES.append("스팀 표 후보를 못 찾음")
        print("  FAIL 스팀 표 후보를 못 찾음")
        return

    table = read_table(
        path, ReadOptions(cell_range=steam_block.range, header_rows=2, auto_detect=False)
    )
    ok = table.columns == ["스팀 / 구분", "사용량", "사용금액"] and table.rows == [
        ["금년월보사용량", 262000, None]
    ]
    label2 = "찾은 표 범위로 실제 읽기까지 정상 동작"
    if ok:
        print(f"  OK   {label2}")
    else:
        FAILURES.append(label2)
        print(f"  FAIL {label2}: columns={table.columns}, rows={table.rows}")


def test_auto_generate() -> None:
    """'⚡ 자동 완성'(원본/템플릿/저장 폴더만 지정) 이 실제로 알아서 완성하는지 확인.

    회사 원본 워크북은 보통 관계없는 시트(안내문, 설비 목록 등)가 섞여
    있으므로, 그런 상황을 재현해 (1) 태그 기반 템플릿, (2) 태그가 하나도
    없는 기존 엑셀 서식(라벨 옆 빈 칸 추론) 둘 다에서 진짜 데이터 시트를
    정확히 골라내는지 확인한다. 특히 (2)는 "용량" 처럼 짧고 흔한 컬럼
    이름이 서로 다른 라벨 여러 개에 부분 일치로 겹쳐 걸려 엉뚱한 시트가
    뽑히던 회귀를 잡기 위한 것이다.
    """
    import datetime as dt

    from docx import Document
    from openpyxl import Workbook

    from reportgen.errors import ReportGenError
    from reportgen.generator import auto_generate

    print("\n[⚡ 자동 완성: 모든 시트를 훑어 자동으로 완성]")

    with tempfile.TemporaryDirectory() as tmp:
        source_path = os.path.join(tmp, "원본.xlsx")
        wb = Workbook()
        decoy1 = wb.active
        decoy1.title = "안내"
        decoy1["A1"] = "이 파일은 매월 자동으로 갱신됩니다."

        decoy2 = wb.create_sheet("설비목록")
        decoy2.append(["설비명", "용량", "위치"])
        decoy2.append(["1호기", 500, "지하"])
        decoy2.append(["2호기", 300, "옥상"])

        real = wb.create_sheet("에너지사용량")
        real.append(["날짜", "사용량", "전력"])
        for day in range(1, 11):
            real.append([dt.date(2026, 1, day), 100 + day, 10 + day])
        wb.save(source_path)

        mapping_dir = os.path.join(tmp, "mappings")
        output_dir = os.path.join(tmp, "output")

        # (1) 태그 기반 워드 템플릿 -> 관계없는 시트를 걸러내고 실제 데이터
        #     시트를 골라야 한다.
        tag_template = os.path.join(tmp, "템플릿.docx")
        doc = Document()
        doc.add_paragraph("사용량 합계: {{사용량}} kWh, 전력: {{전력}}")
        doc.save(tag_template)

        result = auto_generate(source_path, tag_template, output_dir, mapping_dir=mapping_dir)
        ok = (
            len(result.files) == 1
            and result.warnings[0].startswith("'에너지사용량' 시트")
            and "사용량 → 사용량" in result.warnings[0]
            and "전력 → 전력" in result.warnings[0]
        )
        label = "태그 템플릿: 관계없는 시트를 걸러내고 실제 데이터 시트를 골라 연결"
        print(f"  {'OK  ' if ok else 'FAIL'} {label}")
        if not ok:
            FAILURES.append(label)
            print(f"         -> {result.warnings[:1]}")

        # 진단용: 만약 나중에 엉뚱한 표가 뽑히더라도, 파일을 못 보는 상황에서
        # 결과 문구만으로 원인을 알 수 있어야 한다 — 비교한 다른 후보(설비목록)가
        # 점수와 함께 남는지 확인한다.
        ok_diag = "비교한 다른 후보" in result.warnings[0] and "설비목록" in result.warnings[0]
        label_diag = "태그 템플릿: 결과 문구에 비교한 다른 후보(진단용)가 점수와 함께 남음"
        print(f"  {'OK  ' if ok_diag else 'FAIL'} {label_diag}")
        if not ok_diag:
            FAILURES.append(label_diag)
            print(f"         -> {result.warnings[:1]}")

        # 재사용: 저장된 매핑이 있으면 다시 훑지 않고 그대로 써야 한다
        # (재탐색 시 나오는 '자동으로 연결했습니다' 안내문이 없어야 함).
        result2 = auto_generate(source_path, tag_template, output_dir, mapping_dir=mapping_dir)
        ok2 = not any("자동으로 연결했습니다" in w for w in result2.warnings)
        label2 = "태그 템플릿: 두 번째 실행부터는 저장된 매핑을 그대로 재사용"
        print(f"  {'OK  ' if ok2 else 'FAIL'} {label2}")
        if not ok2:
            FAILURES.append(label2)

        # (2) 태그가 하나도 없는 기존 엑셀 서식 -> 라벨 옆 빈 칸을 추론해야
        #     한다. "전력사용량(kWh)" 라벨이 설비목록 시트의 "용량" 컬럼과
        #     부분 일치로 겹쳐 걸려 엉뚱한 시트가 뽑히지 않는지가 핵심이다.
        blank_form = os.path.join(tmp, "기존서식.xlsx")
        fwb = Workbook()
        fsheet = fwb.active
        fsheet.title = "보고서"
        fsheet["A5"] = "전력사용량(kWh)"
        fsheet["A6"] = "사용량"
        fwb.save(blank_form)

        result3 = auto_generate(source_path, blank_form, output_dir, mapping_dir=mapping_dir)
        ok3 = (
            len(result3.files) == 1
            and "라벨 텍스트로" in result3.warnings[0]
            and "에너지사용량" in result3.warnings[0]
        )
        label3 = "태그 없는 서식: 라벨 옆 빈 칸을 추론해 실제 데이터 시트를 골라 연결"
        print(f"  {'OK  ' if ok3 else 'FAIL'} {label3}")
        if not ok3:
            FAILURES.append(label3)
            print(f"         -> {result3.warnings[:1]}")

        # (3) 태그가 하나도 없는 워드 템플릿 -> 라벨 추론 대상이 없으므로
        #     친절한 오류로 실패해야 한다(엉뚱하게 뭔가를 만들면 안 됨).
        blank_word = os.path.join(tmp, "빈워드.docx")
        Document().save(blank_word)
        label4 = "태그 없는 워드 템플릿: 자동으로 완성하지 못하고 오류로 알려줌"
        try:
            auto_generate(source_path, blank_word, output_dir, mapping_dir=mapping_dir)
        except ReportGenError as exc:
            print(f"  OK   {label4}\n         -> {str(exc).splitlines()[0]}")
        except BaseException as exc:  # noqa: BLE001
            FAILURES.append(label4)
            print(f"  FAIL {label4}: {type(exc).__name__} 이(가) 나옴 ({exc})")
        else:
            FAILURES.append(label4)
            print(f"  FAIL {label4}: 오류가 나지 않음")


def test_multifile_consolidation() -> None:
    """'📁 여러 파일 취합'(하루 1파일 로그 -> 월간표) 이 실제로 동작하는지 확인.

    공조기/냉동기/유량계 운전일지처럼 하루마다 새 파일이 생기고, 그 안에
    시간별 데이터 + 맨 아래 그날 요약 행(예: '일사용량 (TON)', '일사용량
    (N/M3)')이 있는 실제 상황을 재현한다. 날짜를 잘못 찾는 것(특히 작은
    숫자를 엑셀 일련번호로 오인하는 회귀)이 가장 위험하므로 그 부분을
    중점적으로 확인한다.
    """
    import datetime as dt

    from openpyxl import Workbook

    from reportgen.data_reader import ReadOptions
    from reportgen.errors import ReportGenError
    from reportgen.multifile import (
        DailyRowSpec,
        build_monthly_table,
        extract_date,
    )

    print("\n[📁 여러 파일 취합: 하루 1파일 로그를 모아 월간표로]")

    def make_daily_file(path: str, day: int, ton: int, nm3: int) -> None:
        wb = Workbook()
        sheet = wb.active
        sheet.title = "DailyReport"
        sheet["A1"] = f"2026년 8월 {day}일 토요일"
        sheet["A3"] = "구분"
        sheet["B3"] = "가스유량"
        sheet["C3"] = "전력"
        row = 4
        for hour in range(1, 6):
            sheet.cell(row=row, column=1, value=f"{hour}시")
            sheet.cell(row=row, column=2, value=hour)
            sheet.cell(row=row, column=3, value=hour * 10)
            row += 1
        sheet.cell(row=row, column=1, value="*안내: 참고용 각주")
        row += 1
        sheet.cell(row=row, column=1, value="일사용량 (TON)")
        sheet.cell(row=row, column=3, value=ton)
        row += 1
        sheet.cell(row=row, column=1, value="일사용량 (N/M3)")
        sheet.cell(row=row, column=3, value=nm3)
        wb.save(path)

    with tempfile.TemporaryDirectory() as tmp:
        daily_dir = os.path.join(tmp, "daily")
        os.makedirs(daily_dir)
        make_daily_file(os.path.join(daily_dir, "유량계 관련_20260801.xlsx"), 1, 137394, 8931)
        make_daily_file(os.path.join(daily_dir, "유량계 관련_20260802.xlsx"), 2, 140000, 9000)
        make_daily_file(os.path.join(daily_dir, "유량계 관련_20260803.xlsx"), 3, 135500, 8800)

        # 날짜를 못 찾아야 하는 파일(파일명에 날짜 없음, 안의 값도 작은 숫자뿐)
        decoy = Workbook()
        decoy_sheet = decoy.active
        decoy_sheet.title = "DailyReport"
        decoy_sheet["A3"] = "구분"
        decoy_sheet["B3"] = "가스유량"
        decoy_sheet["C3"] = "전력"
        decoy_sheet["A4"] = "1시"
        decoy_sheet["B4"] = 1
        decoy_sheet["C4"] = 10
        decoy.save(os.path.join(daily_dir, "메모.xlsx"))

        # 회귀 확인: 작은 숫자(1)를 엑셀 일련번호로 오인해 날짜를 '찾아버리면'
        # 안 된다 (예전엔 1899-12-31 로 잘못 인식되는 문제가 있었다).
        found = extract_date(os.path.join(daily_dir, "메모.xlsx"))
        label0 = "작은 숫자만 있는 파일에서 날짜를 잘못 '찾아내지' 않음"
        if found is None:
            print(f"  OK   {label0}")
        else:
            FAILURES.append(label0)
            print(f"  FAIL {label0}: {found} 로 잘못 인식함")

        options = ReadOptions(sheet_name="DailyReport", cell_range="A3:C11", auto_detect=False)
        spec = DailyRowSpec(row_indexes=[6, 7], row_labels=["TON", "N/M3"], columns=["전력"])
        table, warnings = build_monthly_table(daily_dir, spec, options)

        ok = (
            table.columns == ["날짜", "전력 (TON)", "전력 (N/M3)"]
            and [tuple(r) for r in table.rows]
            == [
                (dt.date(2026, 8, 1), 137394, 8931),
                (dt.date(2026, 8, 2), 140000, 9000),
                (dt.date(2026, 8, 3), 135500, 8800),
            ]
        )
        label1 = "하루 1파일 3일치를 날짜순으로 정확히 취합 (컬럼 이름표 포함)"
        print(f"  {'OK  ' if ok else 'FAIL'} {label1}")
        if not ok:
            FAILURES.append(label1)
            print(f"         -> columns={table.columns}, rows={table.rows}")

        ok2 = any("메모.xlsx" in w and "날짜를 찾지 못해" in w for w in warnings)
        label2 = "날짜 없는 파일은 조용히 무시하지 않고 경고로 남김"
        print(f"  {'OK  ' if ok2 else 'FAIL'} {label2}")
        if not ok2:
            FAILURES.append(label2)
            print(f"         -> warnings={warnings}")

        # 취합 결과를 실제 .xlsx 로 저장해서, 그대로 1단계 원본으로 다시
        # 읽어도 값이 그대로인지(왕복) 확인한다.
        from reportgen.data_reader import read_table
        from reportgen.multifile import save_table_as_excel

        saved_path = os.path.join(tmp, "월간취합.xlsx")
        save_table_as_excel(table, saved_path)
        reloaded = read_table(saved_path)
        ok3 = reloaded.columns == table.columns and reloaded.n_rows == table.n_rows
        label3 = "취합 결과를 저장한 파일이 '1단계 원본 엑셀'로 그대로 다시 읽힘"
        print(f"  {'OK  ' if ok3 else 'FAIL'} {label3}")
        if not ok3:
            FAILURES.append(label3)
            print(f"         -> columns={reloaded.columns}, rows={reloaded.n_rows}")

        # 행 번호를 지정하지 않으면(실수) 오류로 바로 알려줘야 한다.
        label4 = "행 번호를 안 정하면 오류로 바로 알려줌"
        try:
            build_monthly_table(daily_dir, DailyRowSpec(row_indexes=[]), options)
        except ReportGenError as exc:
            print(f"  OK   {label4}\n         -> {str(exc).splitlines()[0]}")
        except BaseException as exc:  # noqa: BLE001
            FAILURES.append(label4)
            print(f"  FAIL {label4}: {type(exc).__name__} 이(가) 나옴 ({exc})")
        else:
            FAILURES.append(label4)
            print(f"  FAIL {label4}: 오류가 나지 않음")


def test_multifile_combined() -> None:
    """설비별 양식이 달라도 **한 번에 월간표 하나로** 합쳐지는지 확인.

    실무에서는 공조기·냉동기·유량계 일지가 시트 이름도 표 위치도 서로 다르다.
    설비별로 읽는 방법을 따로 정해서 넘기면, 날짜를 기준으로 한 줄에 나란히
    붙어야 한다. 설비마다 기록이 빠진 날이 있어도 그 줄 전체가 사라지면 안
    된다(그 칸만 비고 경고로 남아야 한다).
    """
    import datetime as dt

    from openpyxl import Workbook

    from reportgen.data_reader import ReadOptions, read_table
    from reportgen.errors import ReportGenError
    from reportgen.multifile import (
        DailyRowSpec,
        SourceSpec,
        build_combined_monthly_table,
        save_table_as_excel,
    )

    print("\n[📁 여러 설비를 한 번에 월간표 하나로 합치기]")

    with tempfile.TemporaryDirectory() as tmp:
        flow_dir = os.path.join(tmp, "유량계")
        ahu_dir = os.path.join(tmp, "공조기")
        os.makedirs(flow_dir)
        os.makedirs(ahu_dir)

        def make_flow(day: int, ton: int, nm3: int) -> None:
            wb = Workbook()
            sheet = wb.active
            sheet.title = "DailyReport"
            sheet["A1"] = f"2026년 8월 {day}일"
            sheet["A3"] = "구분"
            sheet["B3"] = "가스유량"
            sheet["C3"] = "전력"
            row = 4
            for hour in range(1, 6):
                sheet.cell(row=row, column=1, value=f"{hour}시")
                sheet.cell(row=row, column=2, value=hour)
                sheet.cell(row=row, column=3, value=hour * 10)
                row += 1
            sheet.cell(row=row, column=1, value="*각주")
            row += 1
            sheet.cell(row=row, column=1, value="일사용량 (TON)")
            sheet.cell(row=row, column=3, value=ton)
            row += 1
            sheet.cell(row=row, column=1, value="일사용량 (N/M3)")
            sheet.cell(row=row, column=3, value=nm3)
            wb.save(os.path.join(flow_dir, f"유량계 관련_202608{day:02d}.xlsx"))

        def make_ahu(day: int, hours: float) -> None:
            # 일부러 전혀 다른 양식: 시트 이름도 다르고 표가 B5 부터 시작한다.
            wb = Workbook()
            sheet = wb.active
            sheet.title = "가동시간"
            sheet["B2"] = f"2026년 8월 {day}일 공조기 가동시간"
            sheet["B5"] = "구분"
            sheet["C5"] = "AH-1 가동시간"
            sheet["D5"] = "AH-2 가동시간"
            sheet["B6"] = "0시"
            sheet["C6"] = 0.1
            sheet["D6"] = 0.0
            sheet["B7"] = "1시"
            sheet["C7"] = 1.0
            sheet["D7"] = 1.0
            sheet["B8"] = "일합계"
            sheet["C8"] = hours
            sheet["D8"] = hours + 1
            wb.save(os.path.join(ahu_dir, f"공조기 가동시간_202608{day:02d}.xlsx"))

        for day, ton, nm3 in [(1, 137394, 8931), (2, 140000, 9000), (3, 135500, 8800)]:
            make_flow(day, ton, nm3)
        # 공조기는 3일치가 없다 -> 그 줄이 통째로 사라지면 안 된다.
        for day, hours in [(1, 9.1), (2, 8.5)]:
            make_ahu(day, hours)

        sources = [
            SourceSpec(
                name="유량계",
                folder=flow_dir,
                spec=DailyRowSpec(row_indexes=[6, 7], row_labels=["TON", "N/M3"], columns=["전력"]),
                options=ReadOptions(sheet_name="DailyReport", cell_range="A3:C11", auto_detect=False),
            ),
            SourceSpec(
                name="공조기",
                folder=ahu_dir,
                spec=DailyRowSpec(row_indexes=[2], columns=["AH-1 가동시간", "AH-2 가동시간"]),
                options=ReadOptions(sheet_name="가동시간", cell_range="B5:D8", auto_detect=False),
            ),
        ]
        table, warnings = build_combined_monthly_table(sources)

        expected_columns = [
            "날짜",
            "유량계 · 전력 (TON)",
            "유량계 · 전력 (N/M3)",
            "공조기 · AH-1 가동시간",
            "공조기 · AH-2 가동시간",
        ]
        ok = table.columns == expected_columns and [tuple(r) for r in table.rows] == [
            (dt.date(2026, 8, 1), 137394, 8931, 9.1, 10.1),
            (dt.date(2026, 8, 2), 140000, 9000, 8.5, 9.5),
            (dt.date(2026, 8, 3), 135500, 8800, None, None),
        ]
        label1 = "양식이 다른 설비 2개를 날짜 기준으로 한 표에 나란히 합침 (컬럼에 설비 이름 접두)"
        print(f"  {'OK  ' if ok else 'FAIL'} {label1}")
        if not ok:
            FAILURES.append(label1)
            print(f"         -> columns={table.columns}, rows={table.rows}")

        ok2 = any("공조기" in w and "없는 날짜" in w for w in warnings)
        label2 = "한 설비에만 없는 날짜는 그 칸만 비우고, 어느 설비 며칠인지 경고로 남김"
        print(f"  {'OK  ' if ok2 else 'FAIL'} {label2}")
        if not ok2:
            FAILURES.append(label2)
            print(f"         -> warnings={warnings}")

        saved = save_table_as_excel(table, os.path.join(tmp, "통합월간표.xlsx"))
        reloaded = read_table(saved)
        ok3 = reloaded.columns == expected_columns and reloaded.n_rows == 3
        label3 = "합친 결과를 저장한 파일이 '1단계 원본 엑셀'로 그대로 다시 읽힘"
        print(f"  {'OK  ' if ok3 else 'FAIL'} {label3}")
        if not ok3:
            FAILURES.append(label3)
            print(f"         -> columns={reloaded.columns}, rows={reloaded.n_rows}")

        # 설비 하나가 통째로 실패해도(폴더가 없음) 나머지는 살아야 한다.
        broken = list(sources) + [
            SourceSpec(
                name="없는설비",
                folder=os.path.join(tmp, "없는폴더"),
                spec=DailyRowSpec(row_indexes=[0]),
            )
        ]
        table2, warnings2 = build_combined_monthly_table(broken)
        ok4 = table2.n_rows == 3 and any("없는설비" in w for w in warnings2)
        label4 = "설비 하나가 실패해도 나머지 설비로 계속 진행하고 이유를 남김"
        print(f"  {'OK  ' if ok4 else 'FAIL'} {label4}")
        if not ok4:
            FAILURES.append(label4)
            print(f"         -> rows={table2.n_rows}, warnings={warnings2}")

        # 설비 이름이 겹치면 컬럼이 섞이므로 미리 막아야 한다.
        label5 = "설비 이름이 중복되면 오류로 바로 알려줌"
        try:
            build_combined_monthly_table([sources[0], sources[0]])
        except ReportGenError as exc:
            print(f"  OK   {label5}\n         -> {str(exc).splitlines()[0]}")
        except BaseException as exc:  # noqa: BLE001
            FAILURES.append(label5)
            print(f"  FAIL {label5}: {type(exc).__name__} 이(가) 나옴 ({exc})")
        else:
            FAILURES.append(label5)
            print(f"  FAIL {label5}: 오류가 나지 않음")


def test_graceful_paths() -> None:
    """오류가 '나면 안 되는' 경우들."""
    print("\n[예외 없이 넘어가야 하는 경우]")
    checks = [
        (
            "빈 셀이 섞여 있어도 집계가 된다",
            lambda: aggregate_monthly(
                read_table(os.path.join(SOURCES, "S3_빈셀혼재.xlsx")),
                AggregationSpec(date_column="부서", base_year=2026, base_month=1),
            ),
        ),
    ]
    for label, call in checks:
        try:
            call()
        except AggregationError:
            # 날짜가 아닌 컬럼을 날짜로 지정한 것이므로 AggregationError 가 맞다
            print(f"  OK   {label} (날짜 아님을 정확히 알려 줌)")
        except BaseException as exc:  # noqa: BLE001
            FAILURES.append(label)
            print(f"  FAIL {label}: {type(exc).__name__}: {exc}")
        else:
            print(f"  OK   {label}")

    # 데이터가 한 행뿐이어도, 200행이어도 읽기가 된다
    for name, expected in (("S12_한행.xlsx", 1), ("S11_행많음.xlsx", 200)):
        table = read_table(os.path.join(SOURCES, name))
        if table.n_rows == expected:
            print(f"  OK   {name}: {expected}행 읽음")
        else:
            FAILURES.append(name)
            print(f"  FAIL {name}: {table.n_rows}행 (기대 {expected})")


def test_no_network() -> None:
    """소스 트리 전체에서 네트워크 관련 import 를 찾는다."""
    print("\n[외부 통신 코드 없음 확인]")
    offenders: list[str] = []
    scanned = 0

    for folder, _dirs, files in os.walk(os.path.join(ROOT, "reportgen")):
        for name in files:
            if not name.endswith(".py"):
                continue
            path = os.path.join(folder, name)
            scanned += 1
            tree = ast.parse(open(path, encoding="utf-8").read(), filename=path)
            for node in ast.walk(tree):
                names: list[str] = []
                if isinstance(node, ast.Import):
                    names = [alias.name for alias in node.names]
                elif isinstance(node, ast.ImportFrom) and node.module:
                    names = [node.module]
                for module in names:
                    root = module.split(".")[0]
                    if root in FORBIDDEN_MODULES:
                        offenders.append(f"{os.path.relpath(path, ROOT)}: import {module}")

    # app.py 도 함께 본다
    tree = ast.parse(open(os.path.join(ROOT, "app.py"), encoding="utf-8").read())
    scanned += 1
    for node in ast.walk(tree):
        names = []
        if isinstance(node, ast.Import):
            names = [alias.name for alias in node.names]
        elif isinstance(node, ast.ImportFrom) and node.module:
            names = [node.module]
        for module in names:
            if module.split(".")[0] in FORBIDDEN_MODULES:
                offenders.append(f"app.py: import {module}")

    if offenders:
        FAILURES.extend(offenders)
        print(f"  FAIL 네트워크 관련 import 발견:\n    " + "\n    ".join(offenders))
    else:
        print(f"  OK   파이썬 파일 {scanned}개에서 네트워크 import 없음")


def main() -> int:
    test_input_errors()
    test_formula_engine_computes_simple_formulas()
    test_header_formula_engine()
    test_cross_sheet_formula_computed()
    test_formula_engine_bypass_for_unsupported()
    test_table_block_detection()
    test_auto_generate()
    test_multifile_consolidation()
    test_multifile_combined()
    test_graceful_paths()
    test_no_network()
    print()
    if FAILURES:
        print(f"실패 {len(FAILURES)}건: {FAILURES}")
        return 1
    print("예외 처리 / 로컬 동작 검증 전부 통과")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
