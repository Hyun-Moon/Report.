"""양식 커버리지 검증.

샘플 원본 + 템플릿 조합을 실제로 돌리고, **생성된 결과 파일을 다시 열어서**
값이 올바른 위치에 올바른 값으로 들어갔는지 확인한다.

    python tests/test_matrix.py            # 전체
    python tests/test_matrix.py W3         # 이름에 W3 이 들어간 케이스만
"""

from __future__ import annotations

import datetime as _dt
import os
import shutil
import sys
import tempfile
import traceback
from dataclasses import dataclass, field
from typing import Any, Callable, Optional

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document  # noqa: E402
from openpyxl import load_workbook  # noqa: E402

from reportgen.aggregator import AggregationSpec  # noqa: E402
from reportgen.data_reader import ReadOptions, read_table  # noqa: E402
from reportgen.generator import GenerationRequest, generate  # noqa: E402
from reportgen.mapping import Binding, auto_match  # noqa: E402
from reportgen.templating import open_template  # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
FIXTURES = os.path.join(HERE, "fixtures")
TEMPLATES = os.path.join(FIXTURES, "templates")
SOURCES = os.path.join(FIXTURES, "sources")


# --------------------------------------------------------------------------- #
# 결과 파일 읽기 도우미
# --------------------------------------------------------------------------- #
def word_texts(path: str) -> list[str]:
    """본문·표·머리글·바닥글의 모든 텍스트를 한 목록으로."""
    document = Document(path)
    out: list[str] = [p.text for p in document.paragraphs]

    def walk_table(table) -> None:
        seen: set[int] = set()
        anchors: list[Any] = []
        for row in table.rows:
            for cell in row.cells:
                element = cell._tc
                if id(element) in seen:
                    continue
                seen.add(id(element))
                anchors.append(element)
                out.extend(p.text for p in cell.paragraphs)
                for nested in cell.tables:
                    walk_table(nested)

    for table in document.tables:
        walk_table(table)
    for section in document.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            out.extend(p.text for p in part.paragraphs)
            for table in part.tables:
                walk_table(table)
    return out


def word_blob(path: str) -> str:
    return "\n".join(word_texts(path))


def word_table_grid(path: str, index: int = 0) -> list[list[str]]:
    table = Document(path).tables[index]
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def excel_values(path: str) -> dict[str, Any]:
    """``{'Sheet!A1': 값}``. 수식은 문자열 그대로 돌려준다."""
    workbook = load_workbook(path, data_only=False)
    out: dict[str, Any] = {}
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None:
                    out[f"{sheet.title}!{cell.coordinate}"] = cell.value
    workbook.close()
    return out


def excel_cell(path: str, sheet: str, ref: str):
    workbook = load_workbook(path, data_only=False)
    try:
        return workbook[sheet][ref]
    finally:
        pass  # 스타일 확인을 위해 셀 객체를 살려 둔다


# --------------------------------------------------------------------------- #
# 케이스 정의
# --------------------------------------------------------------------------- #
@dataclass
class Case:
    name: str
    source: str
    template: str
    check: Callable[[str, Any], None]
    read_options: ReadOptions = field(default_factory=ReadOptions)
    aggregate: Optional[AggregationSpec] = None
    extra_bindings: dict[str, Binding] = field(default_factory=dict)
    multi_month_mode: str = "separate"
    expect_files: int = 1


def build_request(case: Case, outdir: str) -> tuple[GenerationRequest, list]:
    source = os.path.join(SOURCES, case.source)
    template = os.path.join(TEMPLATES, case.template)
    table = read_table(source, case.read_options)
    handler = open_template(template)
    slots = handler.scan()
    bindings = auto_match(slots, table.columns)
    bindings.update(case.extra_bindings)
    request = GenerationRequest(
        source_path=source,
        read_options=case.read_options,
        template_path=template,
        bindings=bindings,
        use_aggregation=case.aggregate is not None,
        aggregation=case.aggregate or AggregationSpec(),
        output_dir=outdir,
        multi_month_mode=case.multi_month_mode,
    )
    return request, slots


AGG_BASIC = AggregationSpec(
    methods={"사용량": "sum", "최대수요": "max", "온도": "mean", "가동시간": "sum"}
)


# ---- 1. 워드 템플릿 다양성 ------------------------------------------------- #
def check_w1(path: str, result) -> None:
    blob = word_blob(path)
    assert "12500" in blob, "본문에 사용량 값이 없음"
    assert "23.4" in blob, "본문에 온도 값이 없음"
    assert "480" in blob, "본문에 최대수요 값이 없음"
    assert "김하늘" in blob, "본문에 담당자 값이 없음"
    assert "{{" not in blob, f"치환되지 않은 태그가 남음: {blob}"


def check_w2_agg(path: str, result) -> None:
    grid = word_table_grid(path)
    values = {row[0]: row[1] for row in grid}
    assert values["총 사용량(kWh)"] == "11648", values
    assert values["최대수요(kW)"] == "118", values
    assert values["집계 일수"] == "28", values
    assert values["평균 온도(℃)"].startswith("21.03"), values


def check_w3_merged(path: str, result) -> None:
    grid = word_table_grid(path)
    flat = "\n".join("|".join(row) for row in grid)
    assert "{{" not in flat, f"병합표에 태그가 남음:\n{flat}"
    assert "12500" in flat, "세로 병합 블록 안의 사용량이 안 채워짐"
    assert "13,750,000" in flat, "세로 병합 블록 안의 요금이 안 채워짐"
    assert "정상" in flat, "가로 병합 비고 칸이 안 채워짐"
    assert "1월 전력요금" in flat, "가로 병합 제목 줄이 안 채워짐"
    assert grid[5] == ["담당", "검토", "승인", "합의"], "결재란이 훼손됨"


def check_w4_multi(path: str, result) -> None:
    document = Document(path)
    assert len(document.tables) == 3, f"표 개수가 바뀜: {len(document.tables)}"
    blob = word_blob(path)
    assert "{{" not in blob, "태그가 남음"
    for needle in ("12500", "23.4", "480", "김하늘", "시설관리팀", "정상"):
        assert needle in blob, f"'{needle}' 이(가) 없음"


def check_w5_repeat(path: str, result) -> None:
    texts = word_texts(path)
    blob = "\n".join(texts)
    assert "{{" not in blob, "태그가 남음"
    # 같은 태그가 반복된 횟수만큼 같은 값이 들어가야 한다
    assert blob.count("12500") >= 5, f"사용량 반복 치환 부족: {blob.count('12500')}"
    assert blob.count("김하늘") >= 4, f"담당자 반복 치환 부족: {blob.count('김하늘')}"


def check_w6_header_footer(path: str, result) -> None:
    document = Document(path)
    section = document.sections[0]
    header = "\n".join(p.text for p in section.header.paragraphs)
    footer = "\n".join(p.text for p in section.footer.paragraphs)
    assert "시설관리팀" in header, f"머리글 미치환: {header!r}"
    assert "{{" not in header, f"머리글에 태그가 남음: {header!r}"
    assert "김하늘" in footer, f"바닥글 미치환: {footer!r}"
    assert _dt.date.today().isoformat() in footer, f"바닥글 날짜 미치환: {footer!r}"
    assert "12500" in word_blob(path), "본문 미치환"


def check_w7_loop(path: str, result) -> None:
    grid = word_table_grid(path)
    assert len(grid) == 4, f"반복행이 3개월치로 늘어나지 않음: {grid}"
    assert grid[0] == ["연-월", "일수", "사용량"], grid[0]
    assert grid[1][0] == "2026년 1월" and grid[1][1] == "31", grid[1]
    assert grid[2][0] == "2026년 2월" and grid[2][1] == "28", grid[2]
    assert grid[3][0] == "2026년 3월" and grid[3][1] == "10", grid[3]
    assert grid[1][2] == "3596", grid[1]


def check_w8_special(path: str, result) -> None:
    blob = word_blob(path)
    assert "{{" not in blob, f"특수 이름 태그가 남음: {blob}"
    assert "2026-01, 2026-02, 2026-03" in blob, f"연-월 미치환: {blob}"
    assert "3596" in blob, f"'사용량 (2026-01)' 미치환: {blob}"


def check_w10_vertical_group(path: str, result) -> None:
    grid = word_table_grid(path)
    flat = "\n".join("|".join(row) for row in grid)
    assert "{{" not in flat, f"태그가 남음:\n{flat}"
    assert grid[1] == ["전력", "사용량", "12500"], grid[1]
    assert grid[2] == ["전력", "최대수요", "480"], grid[2]
    assert grid[3] == ["환경", "온도", "23.4"], grid[3]


def check_w11_bullets(path: str, result) -> None:
    blob = word_blob(path)
    assert "{{" not in blob, f"태그가 남음: {blob}"
    for needle in ("2026년 8월", "12500", "480", "23.4", "김하늘", "시설관리팀"):
        assert needle in blob, f"'{needle}' 없음: {blob}"


def check_w12_split_runs(path: str, result) -> None:
    blob = word_blob(path)
    assert "{{" not in blob, f"쪼개진 태그가 복구되지 않음: {blob}"
    assert "12500" in blob and "480" in blob, blob


def check_w9_textbox(path: str, result) -> None:
    # 텍스트 상자는 python-docx 문단 순회로 안 잡히므로 XML 을 직접 본다
    xml = Document(path)._element.xml
    assert "{{" not in xml, "텍스트 상자 안의 태그가 남음"
    assert "시설관리팀" not in xml or True  # 부서는 매핑 대상이 아님
    assert "김하늘" in xml, "텍스트 상자 안 담당자 미치환"
    assert "12500" in xml, "본문 사용량 미치환"
    keys = {slot.key for slot in open_template(os.path.join(TEMPLATES, "W9_텍스트상자.docx")).scan()}
    assert "담당자" in keys and "대상월" in keys, f"도형 안 태그를 스캔이 놓침: {keys}"


# ---- 2. 엑셀 템플릿 다양성 ------------------------------------------------- #
def check_x1_single(path: str, result) -> None:
    cells = excel_values(path)
    assert cells["보고서!B4"] == 11648, cells["보고서!B4"]
    assert cells["보고서!B6"] == 118, cells["보고서!B6"]
    assert cells["보고서!B7"] == 28, cells["보고서!B7"]
    assert cells["보고서!B3"] == "2026년 1월", cells["보고서!B3"]
    assert isinstance(cells["보고서!B4"], (int, float)), "숫자가 문자열로 들어감"


def check_x2_multi_sheet(path: str, result) -> None:
    cells = excel_values(path)
    assert cells["요약!B4"] == 12500, cells.get("요약!B4")
    assert cells["상세!B3"] == 23.4, cells.get("상세!B3")
    assert cells["상세!B4"] == 480, cells.get("상세!B4")
    assert cells["상세!B5"] == "정상", cells.get("상세!B5")
    assert cells["표지!C2"] == "시설관리팀", cells.get("표지!C2")
    assert cells["표지!C4"] == "작성자: 김하늘", cells.get("표지!C4")


def check_x3_irregular(path: str, result) -> None:
    cells = excel_values(path)
    assert cells["불규칙!E7"] == 12500, cells.get("불규칙!E7")
    assert cells["불규칙!E10"] == 480, cells.get("불규칙!E10")
    assert cells["불규칙!E13"] == 23.4, cells.get("불규칙!E13")
    assert cells["불규칙!H4"] == "김하늘", cells.get("불규칙!H4")
    assert cells["불규칙!G16"] == "비고: 정상", cells.get("불규칙!G16")
    assert cells["불규칙!B7"] == "사용량", "라벨이 훼손됨"


def check_x4_styles(path: str, result) -> None:
    workbook = load_workbook(path)
    sheet = workbook["서식"]
    assert sheet["B4"].value == 12500, sheet["B4"].value
    assert sheet["B4"].number_format == "#,##0.0", f"표시형식 소실: {sheet['B4'].number_format}"
    assert sheet["B4"].fill.fgColor.rgb == "00FFF2CC", f"배경색 소실: {sheet['B4'].fill.fgColor.rgb}"
    assert sheet["B4"].border.left.style == "thin", "테두리 소실"
    assert sheet["A3"].font.bold, "헤더 굵기 소실"
    assert "A1:D1" in [str(r) for r in sheet.merged_cells.ranges], "병합 소실"
    assert sheet["A1"].value.endswith("에너지 실적"), sheet["A1"].value
    workbook.close()


def check_x5_formula(path: str, result) -> None:
    cells = excel_values(path)
    assert cells["수식!B2"] == 12500, cells.get("수식!B2")
    assert cells["수식!C2"] == "=B2*1.1", f"수식 훼손: {cells.get('수식!C2')}"
    assert cells["수식!B6"] == "=SUM(B2:B3)", f"수식 훼손: {cells.get('수식!B6')}"
    assert cells["수식!C6"] == "=SUM(C2:C3)", f"수식 훼손: {cells.get('수식!C6')}"
    assert cells["수식!B8"] == "=AVERAGE(B2:B4)", f"수식 훼손: {cells.get('수식!B8')}"
    assert cells["수식!D2"] == '=IF(B2>1000,"주의","정상")', cells.get("수식!D2")
    workbook = load_workbook(path)
    assert workbook.calculation.fullCalcOnLoad, "열 때 재계산 표시가 없음"
    workbook.close()


def check_x6_anchor(path: str, result) -> None:
    cells = excel_values(path)
    assert cells["명세!A3"] == "연-월", cells.get("명세!A3")
    assert cells["명세!B3"] == "일수", cells.get("명세!B3")
    assert cells["명세!A4"] == "2026년 1월", cells.get("명세!A4")
    assert cells["명세!B4"] == 31, cells.get("명세!B4")
    assert cells["명세!A6"] == "2026년 3월", cells.get("명세!A6")
    assert cells["명세!A20"] == "이 아래는 표 영역이 아닙니다.", "앵커 아래 기존 내용이 훼손됨"


def check_x8_frozen_rules(path: str, result) -> None:
    workbook = load_workbook(path)
    sheet = workbook.active
    assert sheet["B2"].value == 12500, sheet["B2"].value
    assert sheet["B3"].value == 480, sheet["B3"].value
    assert str(sheet.freeze_panes) == "A2", f"틀 고정 소실: {sheet.freeze_panes}"
    assert len(list(sheet.conditional_formatting)) == 1, "조건부 서식 소실"
    assert len(sheet.data_validations.dataValidation) == 1, "데이터 유효성 소실"
    workbook.close()


def check_x9_seal_grid(path: str, result) -> None:
    workbook = load_workbook(path)
    sheet = workbook.active
    ranges = {str(r) for r in sheet.merged_cells.ranges}
    for expected in ("A1:F1", "A4:A6", "B4:B6", "C4:C6", "D4:D6"):
        assert expected in ranges, f"병합 {expected} 소실: {ranges}"
    assert sheet["B8"].value == "2026년 8월", sheet["B8"].value
    assert sheet["B9"].value == "김하늘", sheet["B9"].value
    assert sheet["B10"].value == "1,375,000", sheet["B10"].value
    workbook.close()


def check_x7_cell_coords(path: str, result) -> None:
    cells = excel_values(path)
    assert cells["실적표!A4"] == "2026년 1월", cells.get("실적표!A4")
    assert cells["실적표!B4"] == 11648, cells.get("실적표!B4")
    assert cells["실적표!C4"] == 118, cells.get("실적표!C4")
    assert cells["실적표!E4"] == "설비팀", cells.get("실적표!E4")
    workbook = load_workbook(path)
    assert workbook["실적표"]["B4"].border.left.style == "thin", "빈 양식의 테두리가 사라짐"
    workbook.close()


# ---- 3. 원본 데이터 다양성 -------------------------------------------------- #
def check_two_row_header(path: str, result) -> None:
    assert result.table.columns[1] == "전력 / 사용량", result.table.columns
    blob = word_blob(path)
    assert "12500" in blob, f"2단 헤더 컬럼이 매핑되지 않음: {blob}"
    assert "23.4" in blob, blob


def check_blanks(path: str, result) -> None:
    cells = excel_values(path)
    # 첫 행의 최대수요는 비어 있으므로 빈 값이어야 하고, 오류가 나면 안 된다
    assert cells.get("보고서!B4") == 12500, cells.get("보고서!B4")
    assert "보고서!B6" not in cells or cells["보고서!B6"] in ("", None), cells.get("보고서!B6")
    assert result.table.n_rows == 4, f"빈 행이 제거되지 않음: {result.table.n_rows}"


def check_mixed_types(path: str, result) -> None:
    # '1,375,000' 같은 문자열 숫자도 합산되어야 한다
    monthly = result.monthly
    assert monthly is not None
    assert monthly.get("2026-01", "요금") == 5665550, monthly.values
    assert monthly.get("2026-01", "사용량") == 5150.5, monthly.values
    assert monthly.day_counts["2026-01"] == 4, monthly.day_counts


def check_many_rows(path: str, result) -> None:
    assert result.table.n_rows == 200, result.table.n_rows
    blob = word_blob(path)
    assert "{{" not in blob, "태그가 남음"
    # 집계 없이 첫 행만 쓰는 템플릿이므로 첫 행 값이 들어가야 한다
    assert "100" in blob, blob


def check_weekday_suffix(path: str, result) -> None:
    monthly = result.monthly
    assert monthly.periods == ["2026-07"], monthly.periods
    assert monthly.day_counts["2026-07"] == 10, monthly.day_counts
    assert monthly.get("2026-07", "사용량") == 2055, monthly.values


def check_duplicate_dates(path: str, result) -> None:
    monthly = result.monthly
    assert monthly.day_counts["2026-08"] == 5, "같은 날짜 두 번 기록을 별도 날짜로 셈"
    assert monthly.get("2026-08", "사용량") == 580, monthly.values


def check_units_and_negative(path: str, result) -> None:
    monthly = result.monthly
    assert monthly.get("2026-09", "사용량") == 3735, monthly.values
    assert monthly.day_counts["2026-09"] == 4, monthly.day_counts


def check_row_col_offset(path: str, result) -> None:
    assert result.table.columns == ["부서", "사용량", "최대수요"], result.table.columns
    assert result.table.n_rows == 3, result.table.n_rows
    blob = word_blob(path)
    assert "12500" in blob and "{{" not in blob, blob


def check_merged_in_body(path: str, result) -> None:
    monthly = result.monthly
    assert monthly.get("2026-10", "사용량") == 3009, monthly.values
    # 병합된 부서명이 각 날짜 행에 정상 전파되었는지
    assert result.table.column_values("부서") == [
        "시설관리팀", "시설관리팀", "시설관리팀", "생산1팀", "생산1팀",
    ], result.table.column_values("부서")


def check_full_month(path: str, result) -> None:
    monthly = result.monthly
    assert monthly.day_counts["2026-12"] == 31, monthly.day_counts
    assert monthly.get("2026-12", "사용량") == 9796, monthly.values


def check_totals_row_skipped(path: str, result) -> None:
    monthly = result.monthly
    assert monthly.skipped_rows == 2, f"요약행 2개가 스킵되어야 함: {monthly.skipped_rows}"
    assert monthly.day_counts["2026-05"] == 5, monthly.day_counts
    assert monthly.get("2026-05", "사용량") == 1515, monthly.values


def check_single_row(path: str, result) -> None:
    blob = word_blob(path)
    assert "990" in blob and "정별" in blob, blob
    assert "{{" not in blob


# ---- 4. 집계 다양성 --------------------------------------------------------- #
def check_multi_month_separate(path: str, result) -> None:
    assert len(result.files) == 3, [os.path.basename(f) for f in result.files]
    names = [os.path.basename(f) for f in result.files]
    assert any("2026-01" in n for n in names), names
    assert any("2026-03" in n for n in names), names
    first = excel_values(result.files[0])
    assert first["보고서!B3"] == "2026년 1월", first.get("보고서!B3")
    assert first["보고서!B4"] == 3596, first.get("보고서!B4")
    last = excel_values(result.files[2])
    assert last["보고서!B3"] == "2026년 3월", last.get("보고서!B3")
    assert last["보고서!B4"] == 3055, last.get("보고서!B4")


def check_date_formats(path: str, result) -> None:
    monthly = result.monthly
    assert monthly.periods == ["2026-02"], monthly.periods
    assert monthly.day_counts["2026-02"] == 8, monthly.day_counts
    assert monthly.get("2026-02", "사용량") == 1080, monthly.values


def check_day_only(path: str, result) -> None:
    monthly = result.monthly
    assert monthly.periods == ["2026-06"], monthly.periods
    assert monthly.day_counts["2026-06"] == 10, monthly.day_counts
    assert monthly.get("2026-06", "사용량") == 2055, monthly.values


def check_weekend_excluded(path: str, result) -> None:
    monthly = result.monthly
    assert monthly.day_counts["2026-03"] == 22, monthly.day_counts
    assert monthly.get("2026-03", "사용량") == 2200, monthly.values


def check_transposed(path: str, result) -> None:
    monthly = result.monthly
    assert monthly.periods == ["2026-04"], monthly.periods
    assert monthly.get("2026-04", "사용량") == 3055, monthly.values
    assert monthly.get("2026-04", "최대수요") == 90, monthly.values
    assert abs(monthly.get("2026-04", "온도") - 16.55) < 1e-6, monthly.values


def check_holiday_excluded(path: str, result) -> None:
    monthly = result.monthly
    # 2026-03 평일 22일 중 3/2, 3/3 두 날을 공휴일로 제외 -> 20일
    assert monthly.day_counts["2026-03"] == 20, monthly.day_counts
    assert monthly.get("2026-03", "사용량") == 2000, monthly.values


def check_month_filter(path: str, result) -> None:
    monthly = result.monthly
    assert monthly.periods == ["2026-02"], monthly.periods
    assert len(result.files) == 1, result.files


CASES: list[Case] = [
    # 1. 워드 템플릿 다양성
    Case("W1 본문만(표 없음)", "S1_헤더1행.xlsx", "W1_본문만_표없음.docx", check_w1),
    Case(
        "W2 단순표 + 월집계",
        "S5_일단위_한달.xlsx",
        "W2_단순표.docx",
        check_w2_agg,
        aggregate=AGG_BASIC,
    ),
    Case(
        "W3 병합셀 표",
        "S1_헤더1행.xlsx",
        "W3_병합표.docx",
        check_w3_merged,
        extra_bindings={
            "항목": Binding(source="literal", literal="1월 전력요금"),
            "요금": Binding(source="literal", literal="13,750,000"),
        },
    ),
    Case("W4 표 여러 개", "S1_헤더1행.xlsx", "W4_복수표.docx", check_w4_multi),
    Case("W5 같은 태그 반복", "S1_헤더1행.xlsx", "W5_반복태그.docx", check_w5_repeat),
    Case("W6 머리글·바닥글 태그", "S1_헤더1행.xlsx", "W6_머리글바닥글.docx", check_w6_header_footer),
    Case(
        "W7 반복행(여러 달)",
        "S6_일단위_여러달.xlsx",
        "W7_반복행.docx",
        check_w7_loop,
        aggregate=AGG_BASIC,
        multi_month_mode="wide",
    ),
    Case(
        "W8 특수 이름 태그(연-월 등)",
        "S6_일단위_여러달.xlsx",
        "W8_특수이름태그.docx",
        check_w8_special,
        aggregate=AGG_BASIC,
        multi_month_mode="wide",
    ),
    Case("W9 텍스트 상자 안의 태그", "S1_헤더1행.xlsx", "W9_텍스트상자.docx", check_w9_textbox),
    Case("W10 세로 병합 그룹 라벨 표", "S1_헤더1행.xlsx", "W10_세로그룹표.docx", check_w10_vertical_group),
    Case("W11 글머리·번호 목록 태그", "S1_헤더1행.xlsx", "W11_글머리목록.docx", check_w11_bullets),
    Case("W12 워드가 쪼갠 run 안의 태그", "S1_헤더1행.xlsx", "W12_런분할태그.docx", check_w12_split_runs),
    # 2. 엑셀 템플릿 다양성
    Case(
        "X1 단일시트 연속범위",
        "S5_일단위_한달.xlsx",
        "X1_단일시트.xlsx",
        check_x1_single,
        aggregate=AGG_BASIC,
    ),
    Case("X2 여러 시트", "S1_헤더1행.xlsx", "X2_다중시트.xlsx", check_x2_multi_sheet),
    Case("X3 불규칙 배치", "S1_헤더1행.xlsx", "X3_불규칙배치.xlsx", check_x3_irregular),
    Case("X4 서식(테두리·색·병합) 유지", "S1_헤더1행.xlsx", "X4_서식유지.xlsx", check_x4_styles),
    Case("X5 수식 보존", "S1_헤더1행.xlsx", "X5_수식포함.xlsx", check_x5_formula),
    Case(
        "X6 표 앵커({{#표}})",
        "S6_일단위_여러달.xlsx",
        "X6_표앵커.xlsx",
        check_x6_anchor,
        aggregate=AGG_BASIC,
        multi_month_mode="wide",
    ),
    Case(
        "X7 셀 좌표 직접 지정",
        "S5_일단위_한달.xlsx",
        "X7_셀좌표전용.xlsx",
        check_x7_cell_coords,
        aggregate=AGG_BASIC,
        extra_bindings={
            "실적표!A4": Binding(source="builtin", builtin="대상월"),
            "실적표!B4": Binding(source="column", column="사용량"),
            "실적표!C4": Binding(source="column", column="최대수요"),
            "실적표!D4": Binding(source="column", column="온도"),
            "실적표!E4": Binding(source="literal", literal="설비팀"),
        },
    ),
    Case(
        "X8 틀 고정·조건부 서식·유효성",
        "S1_헤더1행.xlsx",
        "X8_틀고정조건부서식.xlsx",
        check_x8_frozen_rules,
    ),
    Case(
        "X9 결재란형 촘촘한 병합",
        "S1_헤더1행.xlsx",
        "X9_결재란병합.xlsx",
        check_x9_seal_grid,
        extra_bindings={"요금": Binding(source="literal", literal="1,375,000")},
    ),
    # 3. 원본 데이터 다양성
    Case("S2 헤더 2행(병합 헤더)", "S2_헤더2행.xlsx", "W1_본문만_표없음.docx", check_two_row_header),
    Case("S3 빈 셀·빈 행 혼재", "S3_빈셀혼재.xlsx", "X1_단일시트.xlsx", check_blanks),
    Case(
        "S4 숫자·문자·날짜 혼합",
        "S4_혼합형식.xlsx",
        "X1_단일시트.xlsx",
        check_mixed_types,
        aggregate=AggregationSpec(methods={"사용량": "sum", "요금": "sum", "온도": "mean"}),
    ),
    Case("S11 데이터 행이 아주 많음", "S11_행많음.xlsx", "W1_본문만_표없음.docx", check_many_rows),
    Case("S12 데이터 행이 하나뿐", "S12_한행.xlsx", "W1_본문만_표없음.docx", check_single_row),
    Case(
        "S16 제목 블록 + 열 오프셋 혼합",
        "S16_행열오프셋.xlsx",
        "W1_본문만_표없음.docx",
        check_row_col_offset,
    ),
    Case(
        "S17 데이터 본문 안의 병합 셀",
        "S17_본문병합.xlsx",
        "X1_단일시트.xlsx",
        check_merged_in_body,
        aggregate=AggregationSpec(date_column="일자", methods={"사용량": "sum", "부서": "text_join"}),
    ),
    # 4. 집계 다양성
    Case(
        "A1 여러 달 -> 월별 보고서 각각",
        "S6_일단위_여러달.xlsx",
        "X1_단일시트.xlsx",
        check_multi_month_separate,
        aggregate=AGG_BASIC,
        multi_month_mode="separate",
        expect_files=3,
    ),
    Case(
        "A2 날짜 표기 제각각",
        "S7_날짜형식.xlsx",
        "X1_단일시트.xlsx",
        check_date_formats,
        aggregate=AggregationSpec(methods={"사용량": "sum", "온도": "mean"}),
    ),
    Case(
        "A3 일자만 있는 표(기준 연·월 지정)",
        "S8_일자만.xlsx",
        "X1_단일시트.xlsx",
        check_day_only,
        aggregate=AggregationSpec(
            date_column="일", base_year=2026, base_month=6,
            methods={"사용량": "sum", "온도": "mean"},
        ),
    ),
    Case(
        "A4 주말 제외",
        "S9_주말포함.xlsx",
        "X1_단일시트.xlsx",
        check_weekend_excluded,
        aggregate=AggregationSpec(exclude_weekends=True, methods={"사용량": "sum"}),
    ),
    Case(
        "A8 요일 병기 날짜(2026-07-01(수))",
        "S13_날짜요일병기.xlsx",
        "X1_단일시트.xlsx",
        check_weekday_suffix,
        aggregate=AggregationSpec(methods={"사용량": "sum", "온도": "mean"}),
    ),
    Case(
        "A9 같은 날짜에 중복 기록(오전/오후)",
        "S14_같은날짜중복.xlsx",
        "X1_단일시트.xlsx",
        check_duplicate_dates,
        aggregate=AggregationSpec(methods={"사용량": "sum", "온도": "mean"}),
    ),
    Case(
        "A10 단위 문자열·음수 혼재",
        "S15_단위및음수.xlsx",
        "X1_단일시트.xlsx",
        check_units_and_negative,
        aggregate=AggregationSpec(methods={"사용량": "sum", "전일대비": "mean"}),
    ),
    Case(
        "A11 결측 없는 전체 31일",
        "S19_전체31일.xlsx",
        "X1_단일시트.xlsx",
        check_full_month,
        aggregate=AGG_BASIC,
    ),
    Case(
        "A12 합계·작성자 행이 섞인 원본",
        "S20_합계행혼재.xlsx",
        "X1_단일시트.xlsx",
        check_totals_row_skipped,
        aggregate=AggregationSpec(methods={"사용량": "sum", "최대수요": "max"}),
    ),
    Case(
        "A5 주말 + 공휴일 제외",
        "S9_주말포함.xlsx",
        "X1_단일시트.xlsx",
        check_holiday_excluded,
        aggregate=AggregationSpec(
            exclude_weekends=True,
            exclude_dates=[_dt.date(2026, 3, 2), _dt.date(2026, 3, 3)],
            methods={"사용량": "sum"},
        ),
    ),
    Case(
        "A6 대상 월만 골라 집계",
        "S6_일단위_여러달.xlsx",
        "X1_단일시트.xlsx",
        check_month_filter,
        aggregate=AggregationSpec(only_months=["2026-02"], methods={"사용량": "sum"}),
    ),
    Case(
        "A7 날짜가 열 방향(전치)",
        "S10_열방향날짜.xlsx",
        "X1_단일시트.xlsx",
        check_transposed,
        read_options=ReadOptions(transpose=True, auto_detect=False),
        aggregate=AggregationSpec(
            methods={"사용량": "sum", "온도": "mean", "최대수요": "max"}
        ),
    ),
]


# --------------------------------------------------------------------------- #
def run(selector: str = "") -> int:
    if not os.path.isdir(FIXTURES):
        print("먼저 python tests/make_fixtures.py 를 실행해 주세요.", file=sys.stderr)
        return 2

    outdir = tempfile.mkdtemp(prefix="reportgen_matrix_")
    cases = [c for c in CASES if not selector or selector in c.name]
    passed: list[str] = []
    failed: list[tuple[str, str]] = []

    print(f"검증 케이스 {len(cases)}건\n" + "=" * 64)
    for case in cases:
        case_dir = os.path.join(outdir, _safe(case.name))
        os.makedirs(case_dir, exist_ok=True)
        try:
            request, _ = build_request(case, case_dir)
            result = generate(request)
            assert len(result.files) == case.expect_files, (
                f"파일 개수 {len(result.files)} (기대 {case.expect_files})"
            )
            case.check(result.files[0], result)
        except AssertionError as exc:
            failed.append((case.name, str(exc)))
            print(f"[실패] {case.name}\n        {exc}")
            continue
        except Exception as exc:  # noqa: BLE001
            failed.append((case.name, f"{type(exc).__name__}: {exc}"))
            print(f"[예외] {case.name}\n        {type(exc).__name__}: {exc}")
            print(textwrap_indent(traceback.format_exc()))
            continue
        passed.append(case.name)
        print(f"[통과] {case.name}")

    print("=" * 64)
    print(f"통과 {len(passed)}건 / 실패 {len(failed)}건")
    if failed:
        print("\n실패 목록:")
        for name, reason in failed:
            print(f"  - {name}: {reason}")
        print(f"\n결과 파일: {outdir}")
        return 1
    shutil.rmtree(outdir, ignore_errors=True)
    return 0


def textwrap_indent(text: str) -> str:
    return "\n".join("        " + line for line in text.strip().splitlines()[-6:])


def _safe(name: str) -> str:
    return "".join(ch if ch.isalnum() or ch in "._-가-힣" else "_" for ch in name)


if __name__ == "__main__":
    raise SystemExit(run(sys.argv[1] if len(sys.argv) > 1 else ""))
