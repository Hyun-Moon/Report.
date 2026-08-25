"""에너지 관리 업무 문서 검증.

사용자가 주로 다루는 문서가 에너지 사용량/절감 관련이라는 점을 반영해,
그 도메인에서 흔한 구조(다중 에너지원, AMI 15분 데이터, 절감량 산정,
계약전력 판정, 온실가스 배출량, 에너지진단 개선안)로 실제 생성까지 돌려
결과 파일을 다시 열어 확인한다.

    python tests/test_energy_matrix.py
"""

from __future__ import annotations

import os
import shutil
import sys
import tempfile
import traceback
from dataclasses import dataclass, field
from typing import Any, Callable, Optional

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document  # noqa: E402
from openpyxl import load_workbook  # noqa: E402

from reportgen.aggregator import AggregationSpec  # noqa: E402
from reportgen.data_reader import read_table  # noqa: E402
from reportgen.generator import GenerationRequest, generate  # noqa: E402
from reportgen.mapping import Binding, auto_match  # noqa: E402
from reportgen.templating import open_template  # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.join(HERE, "fixtures", "energy")
TEMPLATES = os.path.join(ROOT, "templates")
SOURCES = os.path.join(ROOT, "sources")


def word_blob(path: str) -> str:
    document = Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def word_table_grid(path: str, index: int = 0) -> list[list[str]]:
    table = Document(path).tables[index]
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def excel_values(path: str) -> dict[str, Any]:
    workbook = load_workbook(path, data_only=False)
    out: dict[str, Any] = {}
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None:
                    out[f"{sheet.title}!{cell.coordinate}"] = cell.value
    workbook.close()
    return out


@dataclass
class Case:
    name: str
    source: str
    template: str
    check: Callable[[str, Any], None]
    aggregate: Optional[AggregationSpec] = None
    extra_bindings: dict[str, Binding] = field(default_factory=dict)


def build_request(case: Case, outdir: str) -> GenerationRequest:
    source = os.path.join(SOURCES, case.source)
    template = os.path.join(TEMPLATES, case.template)
    table = read_table(source)
    handler = open_template(template)
    bindings = auto_match(handler.scan(), table.columns)
    bindings.update(case.extra_bindings)
    return GenerationRequest(
        source_path=source,
        template_path=template,
        bindings=bindings,
        use_aggregation=case.aggregate is not None,
        aggregation=case.aggregate or AggregationSpec(),
        output_dir=outdir,
    )


# --------------------------------------------------------------------------- #
def check_e1_usage_report(path: str, result) -> None:
    grid = word_table_grid(path)
    values = {row[0]: row[1] for row in grid[1:]}
    assert values["전력사용량"] == "34818", values
    assert values["가스사용량"] == "2436", values
    assert values["유류사용량"] == "392", values
    assert values["최대수요"] == "450", values
    assert values["평균 외기온도"] == "2.35", values
    blob = word_blob(path)
    assert "집계 일수: 28일" in blob, blob
    assert "{{" not in blob, "치환되지 않은 태그가 남음"


def check_e2_diagnosis_measures(path: str, result) -> None:
    grid = word_table_grid(path)
    assert len(grid) == 6, f"개선안 5건이 반복행으로 안 늘어남: {grid}"
    assert grid[1] == ["LED 조명 교체", "8500000", "2400000", "3.5"], grid[1]
    assert grid[5] == ["압축공기 누기 보수", "1800000", "2600000", "0.7"], grid[5]
    blob = word_blob(path)
    assert "개선안 총 5건" in blob, blob
    assert "{{" not in blob, blob


def check_e3_savings(path: str, result) -> None:
    cells = excel_values(path)
    assert cells["절감량산정!B4"] == 15240, cells.get("절감량산정!B4")
    assert cells["절감량산정!B5"] == 12496.8, cells.get("절감량산정!B5")
    assert cells["절감량산정!B6"] == "=B4-B5", "절감량 수식이 깨짐"
    assert cells["절감량산정!B7"] == "=(B4-B5)/B4", "절감률 수식이 깨짐"
    workbook = load_workbook(path)
    assert workbook["절감량산정"]["B7"].number_format == "0.0%", "퍼센트 표시형식 소실"
    workbook.close()


def check_e4_contract_demand(path: str, result) -> None:
    cells = excel_values(path)
    assert cells["계약전력관리!B3"] == 500, cells.get("계약전력관리!B3")
    assert cells["계약전력관리!B4"] == 505, "이상기온 발생일의 최대수요(505)가 안 잡힘"
    assert cells["계약전력관리!B5"] == "=(B3-B4)/B3", "여유율 수식이 깨짐"
    assert cells["계약전력관리!B6"] == '=IF(B4>B3,"초과","정상")', "판정 수식이 깨짐"


def check_c1_chiller_wide(path: str, result) -> None:
    cells = excel_values(path)
    assert cells["운영현황!B4"] == 263.5, cells.get("운영현황!B4")
    assert cells["운영현황!C4"] == 19716, cells.get("운영현황!C4")
    assert cells["운영현황!B5"] == 217, cells.get("운영현황!B5")
    assert cells["운영현황!B6"] == 201.5, cells.get("운영현황!B6")
    assert cells["운영현황!B7"] == "=SUM(B4:B6)", "합계 수식이 깨짐"
    assert cells["운영현황!C7"] == "=SUM(C4:C6)", "합계 수식이 깨짐"


def check_c2_chiller_time_formats(path: str, result) -> None:
    grid = word_table_grid(path)
    values = {row[0]: row[1] for row in grid[1:]}
    # 소수(8.5) + [h]:mm(timedelta) + h:mm(time) + 'HH:MM' 문자열, 나흘치
    # 8.5 x 4 = 34.0 이 정확히 나와야 시간형식 혼재가 올바르게 처리된 것
    assert values["총 가동시간(h)"] == "34", values
    assert values["총 전력사용량(kWh)"] == "2486", values


def check_c4_offseason_zero(path: str, result) -> None:
    grid = word_table_grid(path)
    values = {row[0]: row[1] for row in grid[1:]}
    assert values["총 가동시간(h)"] == "0", f"비수기 0값이 0이 아닌 다른 값으로 나옴: {values}"
    assert values["총 전력사용량(kWh)"] == "0", values
    assert result.monthly.day_counts["2026-01"] == 28, result.monthly.day_counts


def check_c2_equipment_merged(path: str, result) -> None:
    """세로형식으로 여러 설비가 섞이면 설비 구분 없이 합쳐진다는 것을 확인.

    이건 버그가 아니라 이 프로그램의 설계상 한계다(연-월로만 그룹화하고,
    다른 카테고리 축으로는 나누지 않는다). 검증 목적은 '조용히 틀린 값이
    아니라 예측 가능한 합계가 나온다'는 것을 확인하는 것이다.
    """
    monthly = result.monthly
    # 1호기(8.5h×5)+2호기(7.0h×5)+3호기(6.5h×5) = 110.0
    assert monthly.get("2026-08", "가동시간") == 110, monthly.values
    assert monthly.day_counts["2026-08"] == 5, "날짜 자체는 5일로 정확히 셈"


def check_c5_label_not_summed(path: str, result) -> None:
    """'1호기', '3동 2층' 같은 라벨 컬럼이 실수로 합산되지 않는지 확인."""
    from reportgen.aggregator import suggest_methods, detect_date_column

    table = result.table
    suggested = suggest_methods(table, detect_date_column(table))
    assert suggested["설비명"] != "sum", f"라벨 컬럼이 합산으로 추천됨: {suggested}"
    assert suggested["위치"] != "sum", f"라벨 컬럼이 합산으로 추천됨: {suggested}"
    monthly = result.monthly
    assert monthly.get("2026-07", "가동시간") is not None, monthly.values


def check_e5_emission(path: str, result) -> None:
    cells = excel_values(path)
    assert cells["배출량!B3"] == 28892, cells.get("배출량!B3")
    assert cells["배출량!B4"] == 0.4781, cells.get("배출량!B4")
    assert cells["배출량!B5"] == "=B3/1000*B4", "배출량 계산 수식이 깨짐"


# --------------------------------------------------------------------------- #
CASES: list[Case] = [
    Case(
        "E1 다중 에너지원(전력·가스·유류) 종합 보고서",
        "E1_다중에너지원_일별.xlsx",
        "E_T1_에너지사용량종합보고서.docx",
        check_e1_usage_report,
        aggregate=AggregationSpec(
            methods={
                "전력사용량": "sum",
                "가스사용량": "sum",
                "유류사용량": "sum",
                "최대수요": "max",
                "외기온도": "mean",
            }
        ),
    ),
    Case(
        "E2 15분 간격 AMI 수요 데이터(384행) → 월간 최대수요·전력량",
        "E2_15분간격_AMI수요.xlsx",
        "E_T1_에너지사용량종합보고서.docx",
        lambda path, result: _check_e2(path, result),
        aggregate=AggregationSpec(methods={"유효전력량": "sum", "순시수요": "max"}),
        extra_bindings={
            "전력사용량": Binding(source="column", column="유효전력량"),
            "최대수요": Binding(source="column", column="순시수요"),
        },
    ),
    Case(
        "E3 기준 대비 실적 절감량·절감률(수식)",
        "E3_기준실적비교.xlsx",
        "E_T3_절감량산정표.xlsx",
        check_e3_savings,
        aggregate=AggregationSpec(methods={"기준사용량": "sum", "실적사용량": "sum"}),
    ),
    Case(
        "E4 계약전력 초과 판정(IF 수식)",
        "E4_계약전력_최대수요.xlsx",
        "E_T4_계약전력관리대장.xlsx",
        check_e4_contract_demand,
        aggregate=AggregationSpec(methods={"최대수요": "max"}),
        extra_bindings={"계약전력": Binding(source="literal", literal=500)},
    ),
    Case(
        "E5 온실가스 배출량(사용량×배출계수 수식)",
        "E5_온실가스_사용량.xlsx",
        "E_T5_온실가스배출량.xlsx",
        check_e5_emission,
        aggregate=AggregationSpec(methods={"전력사용량": "sum"}),
        extra_bindings={"배출계수": Binding(source="literal", literal=0.4781)},
    ),
    Case(
        "E6 에너지진단 개선안 목록(행 수 유동적)",
        "E6_에너지진단_개선안.xlsx",
        "E_T2_에너지진단개선안.docx",
        check_e2_diagnosis_measures,
        extra_bindings={
            "부서": Binding(source="literal", literal="OO공장"),
            "건수": Binding(source="literal", literal="5"),
        },
    ),
    Case(
        "C1 대형빌딩 냉동기 3대 설비별 집계 + 합계 수식",
        "C1_냉동기3대_넓은형식.xlsx",
        "C_T1_냉동기운영현황.xlsx",
        check_c1_chiller_wide,
        aggregate=AggregationSpec(
            methods={
                "1호기_가동시간": "sum", "1호기_전력사용량": "sum",
                "2호기_가동시간": "sum", "2호기_전력사용량": "sum",
                "3호기_가동시간": "sum", "3호기_전력사용량": "sum",
            }
        ),
    ),
    Case(
        "C2 세로형식 다중설비(연-월로만 그룹화되어 합쳐짐 확인)",
        "C2_냉동기_세로형식.xlsx",
        "C_T2_냉동기가동_단순보고서.docx",
        check_c2_equipment_merged,
        aggregate=AggregationSpec(methods={"가동시간": "sum", "전력사용량": "sum"}),
    ),
    Case(
        "C3 가동시간 시간형식 혼재([h]:mm/h:mm/'HH:MM')",
        "C3_가동시간_시간형식혼재.xlsx",
        "C_T2_냉동기가동_단순보고서.docx",
        check_c2_chiller_time_formats,
        aggregate=AggregationSpec(methods={"가동시간": "sum", "전력사용량": "sum"}),
    ),
    Case(
        "C4 비수기 냉동기 미가동(전부 0)",
        "C4_비수기_전부0.xlsx",
        "C_T2_냉동기가동_단순보고서.docx",
        check_c4_offseason_zero,
        aggregate=AggregationSpec(methods={"가동시간": "sum", "전력사용량": "sum"}),
    ),
    Case(
        "C5 설비 이름표에 숫자 혼재(1호기/3동 2층)",
        "C5_설비이름표_숫자혼재.xlsx",
        "C_T2_냉동기가동_단순보고서.docx",
        check_c5_label_not_summed,
        aggregate=AggregationSpec(methods={"가동시간": "sum"}),
    ),
]


def _check_e2(path: str, result) -> None:
    grid = word_table_grid(path)
    values = {row[0]: row[1] for row in grid[1:]}
    assert values["전력사용량"] == "10120", values  # sum(15,15,...,45,55 ...)
    assert values["최대수요"] == "220", values  # 09~18시 180kW + 14:30 보너스 40kW
    blob = word_blob(path)
    assert "집계 일수: 4일" in blob, blob


def check_derived_value_not_computed() -> None:
    """'원단위(사용량/생산량)' 같은 파생값은 자동 계산되지 않음을 확인한다.

    - 원본에 없는 컬럼을 매핑하면 명확한 오류가 난다.
    - 미리 계산해 컬럼으로 넣어 둔 원본(E7)은 정상 동작한다.
    - 엑셀 템플릿은 수식으로 그 자리에서 계산할 수 있다(E3~E5 로 이미 검증됨).
    """
    from reportgen.errors import MappingError
    from reportgen.mapping import TemplateSlot, resolve_context

    no_intensity = read_table(os.path.join(SOURCES, "E1_다중에너지원_일별.xlsx"))
    try:
        resolve_context(
            [TemplateSlot(key="원단위")],
            {"원단위": Binding(source="column", column="원단위")},
            no_intensity,
        )
    except MappingError:
        pass
    else:
        raise AssertionError("원본에 없는 파생 컬럼인데 오류 없이 통과함")

    precomputed = read_table(os.path.join(SOURCES, "E7_원단위_사전계산.xlsx"))
    context = resolve_context(
        [TemplateSlot(key="원단위")],
        {"원단위": Binding(source="column", column="원단위")},
        precomputed,
    )
    assert context["원단위"] == 3.99, context


def run(selector: str = "") -> int:
    if not os.path.isdir(TEMPLATES):
        print("먼저 python tests/make_energy_fixtures.py 를 실행해 주세요.", file=sys.stderr)
        return 2

    outdir = tempfile.mkdtemp(prefix="reportgen_energy_")
    cases = [c for c in CASES if not selector or selector in c.name]
    passed: list[str] = []
    failed: list[tuple[str, str]] = []

    print(f"에너지 문서 검증 케이스 {len(cases)}건\n" + "=" * 64)
    for case in cases:
        case_dir = os.path.join(outdir, case.name[:20])
        os.makedirs(case_dir, exist_ok=True)
        try:
            request = build_request(case, case_dir)
            result = generate(request)
            case.check(result.files[0], result)
        except AssertionError as exc:
            failed.append((case.name, str(exc)))
            print(f"[실패] {case.name}\n        {exc}")
            continue
        except Exception as exc:  # noqa: BLE001
            failed.append((case.name, f"{type(exc).__name__}: {exc}"))
            print(f"[예외] {case.name}\n        {type(exc).__name__}: {exc}")
            print("\n".join("        " + line for line in traceback.format_exc().splitlines()[-6:]))
            continue
        passed.append(case.name)
        print(f"[통과] {case.name}")

    if not selector:
        try:
            check_derived_value_not_computed()
            print("[통과] 파생값(원단위 등) 자동 계산 안 됨 확인")
            passed.append("파생값 자동계산 제약 확인")
        except AssertionError as exc:
            failed.append(("파생값 자동계산 제약 확인", str(exc)))
            print(f"[실패] 파생값 자동계산 제약 확인\n        {exc}")

    print("=" * 64)
    print(f"통과 {len(passed)}건 / 실패 {len(failed)}건")
    if failed:
        print(f"\n결과 파일: {outdir}")
        return 1
    shutil.rmtree(outdir, ignore_errors=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(run(sys.argv[1] if len(sys.argv) > 1 else ""))
