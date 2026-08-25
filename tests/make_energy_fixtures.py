"""에너지 관리 업무 문서 검증용 샘플 생성기.

사용자가 실제로 주로 다루는 문서가 에너지 사용량/절감 관련이라고 해서,
그 도메인에서 흔한 문서 구조를 재현했다. 특정 기관 서식을 그대로 베끼지
않고 구조·항목 패턴만 참고해 새로 만든 것이다.

다루는 시나리오
---------------
* 여러 에너지원(전력·가스·유류)이 한 표에 섞여 있고 단위가 서로 다른 경우
* 15분 간격 AMI(원격검침) 수요 데이터처럼 하루에도 수십~수백 행이 쌓이는 경우
* 기준(baseline) 사용량과 실적 사용량을 비교해 절감량을 산정하는 표
* 계약전력 대비 최대수요 초과 여부를 판정하는 관리대장
* 사용량에 배출계수를 곱해 온실가스 배출량을 산정하는 표
* 에너지진단 개선안처럼 행 수가 유동적인 목록형 표

    python tests/make_energy_fixtures.py
"""

from __future__ import annotations

import datetime as _dt
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.join(HERE, "fixtures", "energy")
TEMPLATE_DIR = os.path.join(ROOT, "templates")
SOURCE_DIR = os.path.join(ROOT, "sources")

THIN = Side(style="thin", color="808080")
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
HEAD_FILL = PatternFill("solid", fgColor="DCE6F1")
CENTER = Alignment(horizontal="center", vertical="center")


# =========================================================================== #
# 원본 데이터
# =========================================================================== #
def source_e1_multi_energy() -> str:
    """전력(kWh)·가스(Nm³)·유류(L) 가 한 표에 섞이고 단위가 다른 일별 데이터."""
    wb = Workbook()
    ws = wb.active
    ws.title = "일별사용량"
    ws.append(["일자", "전력사용량", "가스사용량", "유류사용량", "최대수요", "외기온도"])
    for day in range(1, 29):  # 2027-02, 결측 없이 28일
        date = _dt.date(2027, 2, day)
        ws.append(
            [
                date,
                1200 + day * 3,          # kWh
                85 + (day % 5),          # Nm3
                12 + (day % 3) * 2,      # L
                420 + (day % 7) * 5,     # kW (최대수요)
                -2.0 + day * 0.3,        # 겨울철 영하권 외기온도
            ]
        )
    return _save_source(wb, "E1_다중에너지원_일별.xlsx")


def source_e2_ami_15min() -> str:
    """15분 간격 AMI(원격검침) 수요 데이터. 하루 96개씩 나흘치 = 384행."""
    wb = Workbook()
    ws = wb.active
    ws.title = "AMI"
    ws.append(["계측시각", "유효전력량", "순시수요"])
    base = _dt.datetime(2026, 6, 1, 0, 0)
    for step in range(96 * 4):  # 4일치
        moment = base + _dt.timedelta(minutes=15 * step)
        hour = moment.hour
        # 주간(9~18시)에 부하가 높은 전형적인 패턴
        load = 180 if 9 <= hour < 18 else 60
        peak_bonus = 40 if hour == 14 and moment.minute == 30 else 0
        ws.append([moment, round((load + peak_bonus) * 0.25, 2), load + peak_bonus])
    return _save_source(wb, "E2_15분간격_AMI수요.xlsx")


def source_e3_baseline_vs_actual() -> str:
    """절감량 산정을 위한 기준(baseline) 대비 실적 사용량 비교 데이터."""
    wb = Workbook()
    ws = wb.active
    ws.title = "비교"
    ws.append(["일자", "기준사용량", "실적사용량"])
    for day in range(1, 31):
        date = _dt.date(2026, 4, day)
        baseline = 500 + (day % 5) * 4
        # 절감 설비 도입 후라 실적이 기준보다 낮게 나온다
        actual = round(baseline * 0.82, 1)
        ws.append([date, baseline, actual])
    return _save_source(wb, "E3_기준실적비교.xlsx")


def source_e4_contract_demand() -> str:
    """계약전력 대비 최대수요 판정용 일별 최대수요 데이터."""
    wb = Workbook()
    ws = wb.active
    ws.title = "최대수요"
    ws.append(["일자", "최대수요"])
    # 대부분 계약전력(500kW) 이내지만 이상기온 발생일에 초과
    values = [420, 435, 410, 460, 505, 440, 415, 470, 490, 430]
    for day, value in enumerate(values, start=1):
        ws.append([_dt.date(2026, 7, day), value])
    return _save_source(wb, "E4_계약전력_최대수요.xlsx")


def source_e5_emission_usage() -> str:
    """온실가스 배출량 산정용 전력사용량만 있는 표 (배출계수는 템플릿에서 곱함)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "사용량"
    ws.append(["일자", "전력사용량"])
    for day in range(1, 32):
        ws.append([_dt.date(2026, 3, day), 900 + day * 2])
    return _save_source(wb, "E5_온실가스_사용량.xlsx")


def source_e6_diagnosis_measures() -> str:
    """에너지진단 개선안 목록. 날짜 없이 항목이 행으로 나열되는 표."""
    wb = Workbook()
    ws = wb.active
    ws.title = "개선안"
    ws.append(["항목", "투자비", "예상절감액", "회수기간"])
    rows = [
        ("LED 조명 교체", 8500000, 2400000, 3.5),
        ("고효율 인버터 도입", 15000000, 4200000, 3.6),
        ("단열 보강", 6200000, 1100000, 5.6),
        ("공조기 폐열회수", 22000000, 5800000, 3.8),
        ("압축공기 누기 보수", 1800000, 2600000, 0.7),
    ]
    for record in rows:
        ws.append(list(record))
    return _save_source(wb, "E6_에너지진단_개선안.xlsx")


def source_e7_derived_intensity_ready() -> str:
    """원단위(사용량/생산량)를 원본에서 이미 계산해 둔 표.

    이 프로그램은 '사용량 ÷ 생산량' 같은 파생 계산을 자동으로 하지 않는다.
    그래서 원단위가 필요하면 원본 엑셀에 미리 계산된 컬럼으로 넣어 둬야
    한다는 것을 보여주는 대조 케이스다.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "원단위"
    ws.append(["일자", "전력사용량", "생산량", "원단위"])
    for day in range(1, 11):
        usage = 1200 + day * 5
        production = 300 + day * 2
        ws.append([_dt.date(2026, 5, day), usage, production, round(usage / production, 3)])
    return _save_source(wb, "E7_원단위_사전계산.xlsx")


# =========================================================================== #
# 템플릿
# =========================================================================== #
def template_e1_usage_report() -> str:
    """전력·가스·유류를 한 문서에 담는 종합 사용량 보고서 (워드)."""
    doc = Document()
    doc.add_heading("에너지 사용량 종합 보고서", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("대상 기간: {{대상월}}   작성일: {{오늘}}")

    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    rows = [
        ("항목", "값", "단위"),
        ("전력사용량", "{{전력사용량}}", "kWh"),
        ("가스사용량", "{{가스사용량}}", "Nm³"),
        ("유류사용량", "{{유류사용량}}", "L"),
        ("최대수요", "{{최대수요}}", "kW"),
        ("평균 외기온도", "{{외기온도}}", "℃"),
    ]
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    doc.add_paragraph("")
    doc.add_paragraph("집계 일수: {{집계일수}}일")
    return _save_doc(doc, "E_T1_에너지사용량종합보고서.docx")


def template_e2_diagnosis_measures() -> str:
    """에너지진단 개선안을 유동적인 행 수로 나열하는 워드 표."""
    doc = Document()
    doc.add_heading("에너지진단 개선안 목록", level=1)
    doc.add_paragraph("사업장: {{부서}}   진단일: {{오늘}}")

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "투자비(원)"
    table.cell(0, 2).text = "예상절감액(원/년)"
    table.cell(0, 3).text = "회수기간(년)"
    table.cell(1, 0).text = "{%tr for r in rows %}"
    table.cell(2, 0).text = "{{ r[0] }}"
    table.cell(2, 1).text = "{{ r[1] }}"
    table.cell(2, 2).text = "{{ r[2] }}"
    table.cell(2, 3).text = "{{ r[3] }}"
    table.cell(3, 0).text = "{%tr endfor %}"

    doc.add_paragraph("")
    doc.add_paragraph("개선안 총 {{건수}}건")
    return _save_doc(doc, "E_T2_에너지진단개선안.docx")


def template_e3_savings_report() -> str:
    """기준 대비 실적을 비교해 절감량·절감률을 수식으로 계산하는 엑셀 표."""
    wb = Workbook()
    ws = wb.active
    ws.title = "절감량산정"

    ws.merge_cells("A1:D1")
    ws["A1"] = "{{대상월}} 절감량 산정표"
    ws["A1"].font = Font(size=13, bold=True)
    ws["A1"].alignment = CENTER
    ws["A1"].fill = HEAD_FILL

    for index, name in enumerate(["구분", "값", "단위", "비고"], start=1):
        cell = ws.cell(row=3, column=index, value=name)
        cell.font = Font(bold=True)
        cell.fill = HEAD_FILL
        cell.border = BOX

    ws["A4"], ws["B4"], ws["C4"] = "기준사용량", "{{기준사용량}}", "kWh"
    ws["A5"], ws["B5"], ws["C5"] = "실적사용량", "{{실적사용량}}", "kWh"
    ws["A6"], ws["B6"], ws["C6"] = "절감량", "=B4-B5", "kWh"
    ws["A7"], ws["B7"], ws["C7"] = "절감률", "=(B4-B5)/B4", "%"
    ws["B7"].number_format = "0.0%"
    for row in range(4, 8):
        for col in range(1, 4):
            ws.cell(row=row, column=col).border = BOX
    return _save_template(wb, "E_T3_절감량산정표.xlsx")


def template_e4_contract_demand_ledger() -> str:
    """계약전력 대비 최대수요 초과 여부를 IF 수식으로 판정하는 관리대장."""
    wb = Workbook()
    ws = wb.active
    ws.title = "계약전력관리"
    ws.merge_cells("A1:C1")
    ws["A1"] = "계약전력 관리대장"
    ws["A1"].font = Font(size=13, bold=True)
    ws["A1"].alignment = CENTER

    ws["A3"], ws["B3"] = "계약전력(kW)", "{{계약전력}}"
    ws["A4"], ws["B4"] = "당월 최대수요(kW)", "{{최대수요}}"
    ws["A5"], ws["B5"] = "여유율", "=(B3-B4)/B3"
    ws["B5"].number_format = "0.0%"
    ws["A6"], ws["B6"] = "판정", '=IF(B4>B3,"초과","정상")'
    for row in range(3, 7):
        ws.cell(row=row, column=1).border = BOX
        ws.cell(row=row, column=2).border = BOX
    return _save_template(wb, "E_T4_계약전력관리대장.xlsx")


def template_e5_emission_report() -> str:
    """사용량 × 배출계수 로 온실가스 배출량을 계산하는 엑셀 표."""
    wb = Workbook()
    ws = wb.active
    ws.title = "배출량"
    ws.merge_cells("A1:C1")
    ws["A1"] = "{{대상월}} 온실가스 배출량 산정"
    ws["A1"].font = Font(size=13, bold=True)
    ws["A1"].alignment = CENTER

    ws["A3"], ws["B3"], ws["C3"] = "전력사용량(kWh)", "{{전력사용량}}", ""
    ws["A4"], ws["B4"], ws["C4"] = "배출계수(tCO2/MWh)", "{{배출계수}}", ""
    ws["A5"], ws["B5"], ws["C5"] = "배출량(tCO2)", "=B3/1000*B4", ""
    for row in range(3, 6):
        ws.cell(row=row, column=1).border = BOX
        ws.cell(row=row, column=2).border = BOX
    return _save_template(wb, "E_T5_온실가스배출량.xlsx")


# =========================================================================== #
# 실행
# =========================================================================== #
def _save_doc(doc: Document, name: str) -> str:
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    path = os.path.join(TEMPLATE_DIR, name)
    doc.save(path)
    return path


def _save_source(wb: Workbook, name: str) -> str:
    os.makedirs(SOURCE_DIR, exist_ok=True)
    path = os.path.join(SOURCE_DIR, name)
    wb.save(path)
    return path


def _save_template(wb: Workbook, name: str) -> str:
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    path = os.path.join(TEMPLATE_DIR, name)
    wb.save(path)
    return path


SOURCE_BUILDERS = [
    source_e1_multi_energy,
    source_e2_ami_15min,
    source_e3_baseline_vs_actual,
    source_e4_contract_demand,
    source_e5_emission_usage,
    source_e6_diagnosis_measures,
    source_e7_derived_intensity_ready,
]

TEMPLATE_BUILDERS = [
    template_e1_usage_report,
    template_e2_diagnosis_measures,
    template_e3_savings_report,
    template_e4_contract_demand_ledger,
    template_e5_emission_report,
]


def build_all() -> dict[str, list[str]]:
    sources = [builder() for builder in SOURCE_BUILDERS]
    templates = [builder() for builder in TEMPLATE_BUILDERS]
    return {"sources": sources, "templates": templates}


if __name__ == "__main__":
    made = build_all()
    for group, paths in made.items():
        print(f"[{group}] {len(paths)}개")
        for path in paths:
            print("   ", os.path.relpath(path, HERE))
