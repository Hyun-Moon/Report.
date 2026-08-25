"""검증용 샘플 파일 생성기.

실무에서 흔히 보는 양식 구조(지출결의서·실적보고서·회의록·통계표 등)의
**레이아웃 패턴만** 참고해서 새로 만든 것이다. 특정 기관의 문구나 로고는
쓰지 않았다.

    python tests/make_fixtures.py

를 실행하면 ``tests/fixtures/`` 아래에 templates/ 와 sources/ 가 생긴다.
"""

from __future__ import annotations

import datetime as _dt
import os
import random

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

HERE = os.path.dirname(os.path.abspath(__file__))
FIXTURES = os.path.join(HERE, "fixtures")
TEMPLATE_DIR = os.path.join(FIXTURES, "templates")
SOURCE_DIR = os.path.join(FIXTURES, "sources")

THIN = Side(style="thin", color="808080")
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
HEAD_FILL = PatternFill("solid", fgColor="DCE6F1")
WARN_FILL = PatternFill("solid", fgColor="FFF2CC")
CENTER = Alignment(horizontal="center", vertical="center")


# =========================================================================== #
# 워드 템플릿
# =========================================================================== #
def word_1_plain() -> str:
    """표가 전혀 없고 본문 문장 사이에 태그만 들어간 문서 (기안문 형태)."""
    doc = Document()
    title = doc.add_heading("월간 에너지 사용 보고", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("수신: 시설관리팀장")
    doc.add_paragraph("작성일: {{오늘}}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "{{대상월}} 기준 전력 사용 실적을 아래와 같이 보고합니다. "
        "해당 월의 총 사용량은 {{사용량}} kWh 이며, 평균 실내온도는 "
        "{{온도}} ℃ 로 확인되었습니다."
    )
    doc.add_paragraph(
        "최대수요는 {{최대수요}} kW 를 기록하였고, 집계에 반영된 일수는 "
        "{{집계일수}}일입니다."
    )
    doc.add_paragraph("")
    doc.add_paragraph("담당: {{담당자}}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return _save_doc(doc, "W1_본문만_표없음.docx")


def word_2_simple_table() -> str:
    """병합 없는 일반 표 하나 (실적보고서 형태)."""
    doc = Document()
    doc.add_heading("월간 실적 보고서", level=1)
    doc.add_paragraph("대상 기간: {{대상월}}")

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    pairs = [
        ("구분", "값"),
        ("총 사용량(kWh)", "{{사용량}}"),
        ("평균 온도(℃)", "{{온도}}"),
        ("최대수요(kW)", "{{최대수요}}"),
        ("집계 일수", "{{집계일수}}"),
    ]
    for row, (left, right) in zip(table.rows, pairs):
        row.cells[0].text = left
        row.cells[1].text = right
    return _save_doc(doc, "W2_단순표.docx")


def word_3_merged_table() -> str:
    """세로/가로 병합이 섞인 복잡한 표 (지출결의서 형태)."""
    doc = Document()
    doc.add_heading("지 출 결 의 서", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"

    # 1행: 가로 병합된 제목 줄
    table.cell(0, 0).merge(table.cell(0, 3)).text = "결의 내용: {{항목}}"

    # 2행: 라벨/값 4칸
    table.cell(1, 0).text = "결의일자"
    table.cell(1, 1).text = "{{오늘}}"
    table.cell(1, 2).text = "대상월"
    table.cell(1, 3).text = "{{대상월}}"

    # 3~4행: 1열을 세로 병합한 '금액' 블록
    table.cell(2, 0).merge(table.cell(3, 0)).text = "금액"
    table.cell(2, 1).text = "사용량(kWh)"
    table.cell(2, 2).merge(table.cell(2, 3)).text = "{{사용량}}"
    table.cell(3, 1).text = "요금(원)"
    table.cell(3, 2).merge(table.cell(3, 3)).text = "{{요금}}"

    # 5행: 비고
    table.cell(4, 0).text = "비고"
    table.cell(4, 1).merge(table.cell(4, 3)).text = "{{비고}}"

    # 6행: 결재란 (가로 4칸)
    for index, label in enumerate(["담당", "검토", "승인", "합의"]):
        table.cell(5, index).text = label

    doc.add_paragraph("")
    doc.add_paragraph("작성자: {{담당자}}")
    return _save_doc(doc, "W3_병합표.docx")


def word_4_multi_table() -> str:
    """표가 여러 개 섞여 있고 사이에 본문이 끼어 있는 문서 (회의록 형태)."""
    doc = Document()
    doc.add_heading("정기 점검 회의록", level=1)

    doc.add_paragraph("1. 개요")
    head = doc.add_table(rows=2, cols=4)
    head.style = "Table Grid"
    head.cell(0, 0).text = "일시"
    head.cell(0, 1).text = "{{오늘}}"
    head.cell(0, 2).text = "대상월"
    head.cell(0, 3).text = "{{대상월}}"
    head.cell(1, 0).text = "작성자"
    head.cell(1, 1).text = "{{담당자}}"
    head.cell(1, 2).text = "부서"
    head.cell(1, 3).text = "{{부서}}"

    doc.add_paragraph("")
    doc.add_paragraph("2. 실적 요약")
    body = doc.add_table(rows=4, cols=3)
    body.style = "Table Grid"
    rows = [
        ("항목", "값", "단위"),
        ("사용량", "{{사용량}}", "kWh"),
        ("평균온도", "{{온도}}", "℃"),
        ("최대수요", "{{최대수요}}", "kW"),
    ]
    for row, values in zip(body.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    doc.add_paragraph("")
    doc.add_paragraph("3. 특이사항")
    note = doc.add_table(rows=1, cols=1)
    note.style = "Table Grid"
    note.cell(0, 0).text = "{{비고}}"

    doc.add_paragraph("")
    doc.add_paragraph("작성: {{담당자}} / 확인: {{부서}}")
    return _save_doc(doc, "W4_복수표.docx")


def word_5_repeated_tags() -> str:
    """같은 태그가 문서 곳곳에서 반복되는 문서."""
    doc = Document()
    doc.add_heading("{{대상월}} 운영 요약", level=1)
    doc.add_paragraph("{{대상월}} 실적을 정리한 문서입니다. (작성 {{담당자}})")
    doc.add_paragraph("사용량은 {{사용량}} kWh 입니다.")

    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "기간"
    table.cell(0, 1).text = "{{대상월}}"
    table.cell(0, 2).text = "{{담당자}}"
    table.cell(1, 0).text = "사용량"
    table.cell(1, 1).text = "{{사용량}}"
    table.cell(1, 2).text = "{{사용량}}"
    table.cell(2, 0).text = "재확인"
    table.cell(2, 1).text = "{{사용량}}"
    table.cell(2, 2).text = "{{대상월}}"

    doc.add_paragraph("끝으로 {{대상월}} 사용량 {{사용량}} kWh 를 확정합니다.")
    doc.add_paragraph("담당 {{담당자}}, 담당 {{담당자}}")
    return _save_doc(doc, "W5_반복태그.docx")


def word_6_header_footer() -> str:
    """머리글/바닥글에 태그가 들어간 문서."""
    doc = Document()
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.text = "{{부서}} | {{대상월}} 정기보고"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer.paragraphs[0]
    footer.text = "작성자 {{담당자}} · 출력 {{오늘}}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("정기 보고", level=1)
    doc.add_paragraph("본문 사용량: {{사용량}} kWh")

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "최대수요"
    table.cell(0, 1).text = "{{최대수요}}"
    table.cell(1, 0).text = "평균온도"
    table.cell(1, 1).text = "{{온도}}"
    return _save_doc(doc, "W6_머리글바닥글.docx")


def word_7_loop_table() -> str:
    """docxtpl 반복문으로 행 수가 유동적인 표를 만드는 문서 (통계표 형태)."""
    doc = Document()
    doc.add_heading("월별 집계표", level=1)
    doc.add_paragraph("대상: {{대상월}}")

    # docxtpl 의 행 반복은 표식이 들어간 '행 자체'를 통째로 없앤다. 그래서
    # {%tr for %} 와 {%tr endfor %} 는 각각 자기 행을 따로 차지해야 하고,
    # 그 사이에 낀 행이 데이터 수만큼 복제된다.
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "연-월"
    table.cell(0, 1).text = "일수"
    table.cell(0, 2).text = "사용량"
    table.cell(1, 0).text = "{%tr for r in rows %}"
    table.cell(2, 0).text = "{{ r[0] }}"
    table.cell(2, 1).text = "{{ r[1] }}"
    table.cell(2, 2).text = "{{ r[2] }}"
    table.cell(3, 0).text = "{%tr endfor %}"

    doc.add_paragraph("합계 확인용 태그: {{사용량}}")
    return _save_doc(doc, "W7_반복행.docx")


def word_8_hyphen_tag() -> str:
    """jinja2 가 연산자로 오해하는 이름(``연-월``)을 태그로 쓴 문서."""
    doc = Document()
    doc.add_heading("특수 이름 태그 시험", level=1)
    doc.add_paragraph("대상 키: {{연-월}}")
    doc.add_paragraph("사용량(2026-01): {{사용량 (2026-01)}}")
    doc.add_paragraph("띄어쓰기 태그: {{총 사용량}}")
    return _save_doc(doc, "W8_특수이름태그.docx")


def word_9_textbox() -> str:
    """텍스트 상자(도형) 안에 태그가 들어간 문서.

    python-docx 의 문단/표 순회로는 보이지 않는 영역이라, 스캔이 이를 놓치면
    보고서에 조용히 빈칸이 나간다.
    """
    doc = Document()
    doc.add_heading("도형 포함 문서", level=1)
    doc.add_paragraph("본문 사용량: {{사용량}}")

    holder = doc.add_paragraph()
    shape_xml = (
        '<w:r %s>'
        '<mc:AlternateContent'
        ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
        "<mc:Fallback><w:pict>"
        '<v:rect xmlns:v="urn:schemas-microsoft-com:vml" style="width:240pt;height:48pt">'
        "<v:textbox><w:txbxContent><w:p><w:r>"
        "<w:t>요약: {{대상월}} / 담당 {{담당자}}</w:t>"
        "</w:r></w:p></w:txbxContent></v:textbox>"
        "</v:rect></w:pict></mc:Fallback>"
        "</mc:AlternateContent></w:r>" % nsdecls("w")
    )
    holder._p.append(parse_xml(shape_xml))

    doc.add_paragraph("최대수요: {{최대수요}}")
    return _save_doc(doc, "W9_텍스트상자.docx")


def word_10_vertical_group_table() -> str:
    """첫 열이 세로로 병합되어 항목을 묶는 표 (전력/환경 그룹 라벨)."""
    doc = Document()
    doc.add_heading("그룹 라벨 병합 표", level=1)

    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "구분"
    table.cell(0, 1).text = "항목"
    table.cell(0, 2).text = "값"

    table.cell(1, 0).merge(table.cell(2, 0)).text = "전력"
    table.cell(1, 1).text = "사용량"
    table.cell(1, 2).text = "{{사용량}}"
    table.cell(2, 1).text = "최대수요"
    table.cell(2, 2).text = "{{최대수요}}"

    table.cell(3, 0).merge(table.cell(4, 0)).text = "환경"
    table.cell(3, 1).text = "온도"
    table.cell(3, 2).text = "{{온도}}"
    table.cell(4, 1).text = "비고"
    table.cell(4, 2).text = "{{비고}}"
    return _save_doc(doc, "W10_세로그룹표.docx")


def word_11_bulleted_list() -> str:
    """글머리 기호·번호 매기기 목록 안에 태그가 들어간 문서."""
    doc = Document()
    doc.add_heading("점검 결과 목록", level=1)
    doc.add_paragraph(f"대상월: {{{{대상월}}}}")

    for text in (
        f"사용량: {{{{사용량}}}} kWh",
        f"최대수요: {{{{최대수요}}}} kW",
        f"평균온도: {{{{온도}}}} ℃",
    ):
        doc.add_paragraph(text, style="List Bullet")

    for text in (
        f"1차 확인 - 담당 {{{{담당자}}}}",
        f"2차 확인 - 부서 {{{{부서}}}}",
    ):
        doc.add_paragraph(text, style="List Number")
    return _save_doc(doc, "W11_글머리목록.docx")


def word_12_split_runs() -> str:
    """워드가 태그 글자를 여러 서식 조각(run)으로 쪼개 놓은 실제 상황 재현.

    사람이 태그 일부만 볼드로 칠하거나 맞춤법 검사가 끼어들면, 워드는
    ``{{사용량}}`` 을 ``{{`` / ``사용량`` / ``}}`` 처럼 서로 다른 run 으로
    쪼개서 저장한다. docxtpl 의 patch_xml 단계가 이를 복구하는지 확인한다.
    """
    doc = Document()
    doc.add_heading("run 분할 태그 시험", level=1)

    paragraph = doc.add_paragraph()
    paragraph.add_run("사용량은 {{")
    bold_run = paragraph.add_run("사용량")
    bold_run.bold = True
    paragraph.add_run("}} kWh 이고, 최대수요는 {{최대")
    paragraph.add_run("수요}} kW 입니다.")
    return _save_doc(doc, "W12_런분할태그.docx")


# =========================================================================== #
# 엑셀 템플릿
# =========================================================================== #
def excel_1_single_sheet() -> str:
    """단일 시트, 위에서 아래로 연속된 라벨/값 구조."""
    wb = Workbook()
    ws = wb.active
    ws.title = "보고서"
    ws["A1"] = "월간 실적 보고"
    ws["A1"].font = Font(size=14, bold=True)
    rows = [
        ("대상월", "{{대상월}}"),
        ("사용량(kWh)", "{{사용량}}"),
        ("평균온도(℃)", "{{온도}}"),
        ("최대수요(kW)", "{{최대수요}}"),
        ("집계일수", "{{집계일수}}"),
        ("담당자", "{{담당자}}"),
    ]
    for offset, (label, tag) in enumerate(rows, start=3):
        ws.cell(row=offset, column=1, value=label).border = BOX
        ws.cell(row=offset, column=2, value=tag).border = BOX
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 22
    return _save_wb(wb, "X1_단일시트.xlsx")


def excel_2_multi_sheet() -> str:
    """요약 시트와 상세 시트에 나눠 값을 채워야 하는 구조."""
    wb = Workbook()
    summary = wb.active
    summary.title = "요약"
    summary["A1"] = "요약"
    summary["A1"].font = Font(bold=True)
    summary["A3"] = "대상월"
    summary["B3"] = "{{대상월}}"
    summary["A4"] = "총사용량"
    summary["B4"] = "{{사용량}}"

    detail = wb.create_sheet("상세")
    detail["A1"] = "상세 항목"
    detail["A1"].font = Font(bold=True)
    detail["A3"] = "평균온도"
    detail["B3"] = "{{온도}}"
    detail["A4"] = "최대수요"
    detail["B4"] = "{{최대수요}}"
    detail["A5"] = "비고"
    detail["B5"] = "{{비고}}"

    cover = wb.create_sheet("표지")
    cover["C2"] = "{{부서}}"
    cover["C4"] = "작성자: {{담당자}}"
    return _save_wb(wb, "X2_다중시트.xlsx")


def excel_3_irregular() -> str:
    """값 칸이 열/행을 건너뛰며 흩어져 있는 구조."""
    wb = Workbook()
    ws = wb.active
    ws.title = "불규칙"
    ws["B2"] = "보고 개요"
    ws["B2"].font = Font(bold=True)
    ws["B4"] = "대상월"
    ws["D4"] = "{{대상월}}"
    ws["F4"] = "담당"
    ws["H4"] = "{{담당자}}"

    ws["B7"] = "사용량"
    ws["E7"] = "{{사용량}}"
    ws["B10"] = "최대수요"
    ws["E10"] = "{{최대수요}}"
    ws["B13"] = "평균온도"
    ws["E13"] = "{{온도}}"
    ws["G16"] = "비고: {{비고}}"
    return _save_wb(wb, "X3_불규칙배치.xlsx")


def excel_4_styled() -> str:
    """테두리/배경색/병합/표시형식이 살아 있어야 하는 구조."""
    wb = Workbook()
    ws = wb.active
    ws.title = "서식"

    ws.merge_cells("A1:D1")
    ws["A1"] = "{{대상월}} 에너지 실적"
    ws["A1"].font = Font(size=13, bold=True, color="1F4E79")
    ws["A1"].alignment = CENTER
    ws["A1"].fill = HEAD_FILL

    headers = ["구분", "값", "단위", "비고"]
    for index, name in enumerate(headers, start=1):
        cell = ws.cell(row=3, column=index, value=name)
        cell.font = Font(bold=True)
        cell.fill = HEAD_FILL
        cell.border = BOX
        cell.alignment = CENTER

    data = [
        ("사용량", "{{사용량}}", "kWh", ""),
        ("최대수요", "{{최대수요}}", "kW", ""),
        ("평균온도", "{{온도}}", "℃", "{{비고}}"),
    ]
    for r_offset, record in enumerate(data, start=4):
        for c_offset, value in enumerate(record, start=1):
            cell = ws.cell(row=r_offset, column=c_offset, value=value)
            cell.border = BOX
            if c_offset == 2:
                cell.number_format = "#,##0.0"
                cell.fill = WARN_FILL
    return _save_wb(wb, "X4_서식유지.xlsx")


def excel_5_formula() -> str:
    """값을 채우는 셀 바로 옆/아래에 수식이 있는 구조."""
    wb = Workbook()
    ws = wb.active
    ws.title = "수식"
    ws["A1"] = "항목"
    ws["B1"] = "값"
    ws["C1"] = "환산(×1.1)"
    for index, (label, tag) in enumerate(
        [("사용량", "{{사용량}}"), ("최대수요", "{{최대수요}}"), ("평균온도", "{{온도}}")],
        start=2,
    ):
        ws.cell(row=index, column=1, value=label)
        ws.cell(row=index, column=2, value=tag)
        ws.cell(row=index, column=3, value=f"=B{index}*1.1")
    ws["A6"] = "합계"
    ws["B6"] = "=SUM(B2:B3)"
    ws["C6"] = "=SUM(C2:C3)"
    ws["A8"] = "평균"
    ws["B8"] = "=AVERAGE(B2:B4)"
    ws["D2"] = '=IF(B2>1000,"주의","정상")'
    return _save_wb(wb, "X5_수식포함.xlsx")


def excel_6_table_anchor() -> str:
    """``{{#표}}`` 자리부터 집계표를 통째로 써 넣는 구조."""
    wb = Workbook()
    ws = wb.active
    ws.title = "명세"
    ws["A1"] = "{{대상월}} 상세 명세"
    ws["A1"].font = Font(bold=True)
    ws["A3"] = "{{#표}}"
    ws["A20"] = "이 아래는 표 영역이 아닙니다."
    return _save_wb(wb, "X6_표앵커.xlsx")


def excel_7_cell_coords() -> str:
    """태그가 하나도 없고, 셀 좌표를 직접 지정해서 채우는 기존 양식."""
    wb = Workbook()
    ws = wb.active
    ws.title = "실적표"
    ws.merge_cells("A1:E1")
    ws["A1"] = "월간 실적 집계표"
    ws["A1"].alignment = CENTER
    ws["A1"].font = Font(size=13, bold=True)

    for index, name in enumerate(["대상월", "사용량", "최대수요", "평균온도", "담당"], start=1):
        cell = ws.cell(row=3, column=index, value=name)
        cell.fill = HEAD_FILL
        cell.border = BOX
        cell.alignment = CENTER
    for index in range(1, 6):
        ws.cell(row=4, column=index).border = BOX
    return _save_wb(wb, "X7_셀좌표전용.xlsx")


def excel_8_frozen_and_rules() -> str:
    """틀 고정 + 조건부 서식 + 데이터 유효성이 함께 걸린 양식."""
    from openpyxl.formatting.rule import CellIsRule
    from openpyxl.worksheet.datavalidation import DataValidation

    wb = Workbook()
    ws = wb.active
    ws.title = "점검표"
    for index, name in enumerate(["항목", "값", "판정"], start=1):
        cell = ws.cell(row=1, column=index, value=name)
        cell.font = Font(bold=True)
        cell.fill = HEAD_FILL
    ws["A2"], ws["B2"] = "사용량", "{{사용량}}"
    ws["A3"], ws["B3"] = "최대수요", "{{최대수요}}"
    ws["A4"], ws["B4"] = "온도", "{{온도}}"
    ws.freeze_panes = "A2"

    ws.conditional_formatting.add(
        "B2:B4",
        CellIsRule(operator="greaterThan", formula=["1000"], fill=WARN_FILL),
    )
    validation = DataValidation(type="list", formula1='"정상,주의,점검필요"')
    ws.add_data_validation(validation)
    validation.add("C2:C4")
    return _save_wb(wb, "X8_틀고정조건부서식.xlsx")


def excel_9_seal_grid() -> str:
    """결재란처럼 작은 셀들이 촘촘히 병합된 표 + 별도 위치의 태그."""
    wb = Workbook()
    ws = wb.active
    ws.title = "기안문"
    ws.merge_cells("A1:F1")
    ws["A1"] = "지출 기안서"
    ws["A1"].font = Font(size=13, bold=True)
    ws["A1"].alignment = CENTER

    # 결재란: 2행짜리 라벨 + 서명칸이 4명분, 각 칸은 세로 병합
    labels = ["기안", "검토", "협조", "결재"]
    for index, label in enumerate(labels):
        col = index + 1
        ws.merge_cells(start_row=3, start_column=col, end_row=3, end_column=col)
        ws.cell(row=3, column=col, value=label).alignment = CENTER
        ws.cell(row=3, column=col).fill = HEAD_FILL
        ws.merge_cells(start_row=4, start_column=col, end_row=6, end_column=col)
        cell = ws.cell(row=4, column=col)
        cell.border = BOX
        cell.alignment = Alignment(horizontal="center", vertical="top")

    ws["A8"] = "대상월"
    ws["B8"] = "{{대상월}}"
    ws["A9"] = "기안자"
    ws["B9"] = "{{담당자}}"
    ws["A10"] = "금액"
    ws["B10"] = "{{요금}}"
    return _save_wb(wb, "X9_결재란병합.xlsx")


# =========================================================================== #
# 원본 데이터
# =========================================================================== #
def source_1_single_header() -> str:
    """헤더 1행짜리 가장 기본적인 표."""
    wb = Workbook()
    ws = wb.active
    ws.title = "실적"
    ws.append(["부서", "담당자", "사용량", "최대수요", "온도", "비고"])
    ws.append(["시설관리팀", "김하늘", 12500, 480, 23.4, "정상"])
    ws.append(["생산1팀", "이바다", 30800, 910, 24.1, ""])
    ws.append(["생산2팀", "박구름", 21450, 760, 22.8, "설비 점검"])
    return _save_wb(wb, "S1_헤더1행.xlsx", SOURCE_DIR)


def source_2_two_row_header() -> str:
    """상위 헤더가 병합된 2단 헤더 표 (통계표 형태)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "집계"
    ws.merge_cells("A1:A2")
    ws["A1"] = "부서"
    ws.merge_cells("B1:C1")
    ws["B1"] = "전력"
    ws["B2"] = "사용량"
    ws["C2"] = "최대수요"
    ws.merge_cells("D1:E1")
    ws["D1"] = "환경"
    ws["D2"] = "온도"
    ws["E2"] = "습도"
    rows = [
        ("시설관리팀", 12500, 480, 23.4, 51),
        ("생산1팀", 30800, 910, 24.1, 48),
        ("생산2팀", 21450, 760, 22.8, 55),
    ]
    for record in rows:
        ws.append(list(record))
    return _save_wb(wb, "S2_헤더2행.xlsx", SOURCE_DIR)


def source_3_blanks() -> str:
    """빈 셀과 빈 행이 섞여 있고, 표가 시트 중간에서 시작하는 경우."""
    wb = Workbook()
    ws = wb.active
    ws.title = "원본"
    ws["B2"] = "2026년 상반기 집계 자료"  # 표가 아닌 제목 줄
    ws["B4"] = "부서"
    ws["C4"] = "사용량"
    ws["D4"] = "최대수요"
    ws["E4"] = "온도"
    data = [
        ("시설관리팀", 12500, None, 23.4),
        ("생산1팀", None, 910, None),
        (None, None, None, None),  # 빈 행
        ("생산2팀", 21450, 760, 22.8),
        ("연구소", 8000, 300, None),
    ]
    for offset, record in enumerate(data, start=5):
        for c_offset, value in enumerate(record, start=2):
            if value is not None:
                ws.cell(row=offset, column=c_offset, value=value)
    return _save_wb(wb, "S3_빈셀혼재.xlsx", SOURCE_DIR)


def source_4_mixed_types() -> str:
    """숫자/문자/날짜/천단위 콤마 문자열이 섞인 표."""
    wb = Workbook()
    ws = wb.active
    ws.title = "혼합"
    ws.append(["일자", "부서", "사용량", "요금", "온도", "판정"])
    ws.append([_dt.date(2026, 1, 5), "시설관리팀", 1250, "1,375,000", 23.4, "정상"])
    ws.append(["2026-01-06", "시설관리팀", "1,180", 1298000, "22.9", "정상"])
    ws.append([_dt.datetime(2026, 1, 7, 9, 30), "시설관리팀", 1320.5, "1,452,550", 24.0, "주의"])
    ws.append([46030, "시설관리팀", 1400, "1,540,000", 23.1, "정상"])  # 엑셀 일련번호
    return _save_wb(wb, "S4_혼합형식.xlsx", SOURCE_DIR)


def source_5_daily_one_month() -> str:
    """한 달치 일단위 데이터. 중간에 날짜가 빠져 있다 (28일치)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "일별"
    ws.append(["일자", "사용량", "최대수요", "온도", "가동시간"])
    random.seed(20260101)
    day = 1
    written = 0
    while written < 28 and day <= 31:
        if day in (11, 12, 25):  # 결측일
            day += 1
            continue
        ws.append(
            [
                _dt.date(2026, 1, day),
                400 + day,
                100 + (day % 7) * 3,
                20.0 + (day % 5) * 0.5,
                8,
            ]
        )
        written += 1
        day += 1
    return _save_wb(wb, "S5_일단위_한달.xlsx", SOURCE_DIR)


def source_6_daily_multi_month() -> str:
    """1~3월 일단위 데이터가 한 파일에 섞여 있는 경우."""
    wb = Workbook()
    ws = wb.active
    ws.title = "일별"
    ws.append(["일자", "사용량", "최대수요", "온도"])
    for month, days in ((1, 31), (2, 28), (3, 10)):
        for day in range(1, days + 1):
            ws.append(
                [
                    _dt.date(2026, month, day),
                    100 * month + day,
                    50 + day,
                    15.0 + month + (day % 3),
                ]
            )
    return _save_wb(wb, "S6_일단위_여러달.xlsx", SOURCE_DIR)


def source_7_date_formats() -> str:
    """날짜 표기가 제각각인 경우."""
    wb = Workbook()
    ws = wb.active
    ws.title = "일별"
    ws.append(["일자", "사용량", "온도"])
    ws.append(["2026-02-01", 100, 20.0])
    ws.append(["2026/02/02", 110, 21.0])
    ws.append(["2026.02.03", 120, 22.0])
    ws.append(["20260204", 130, 23.0])
    ws.append(["2026년 2월 5일", 140, 24.0])
    ws.append([46059, 150, 25.0])  # 엑셀 일련번호 = 2026-02-06
    ws.append(["02/07", 160, 26.0])
    ws.append([_dt.date(2026, 2, 8), 170, 27.0])
    return _save_wb(wb, "S7_날짜형식.xlsx", SOURCE_DIR)


def source_8_day_only() -> str:
    """'1일, 2일 ...' 처럼 일자만 있는 표 (기준 연-월을 따로 줘야 함)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "일계"
    ws.append(["일", "사용량", "온도"])
    for day in range(1, 11):
        ws.append([f"{day}일", 200 + day, 18.0 + day * 0.2])
    return _save_wb(wb, "S8_일자만.xlsx", SOURCE_DIR)


def source_9_weekend() -> str:
    """주말이 섞여 있어 평일만 집계해야 하는 경우 (2026-03)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "일별"
    ws.append(["일자", "사용량", "구분"])
    for day in range(1, 32):
        date = _dt.date(2026, 3, day)
        weekend = date.weekday() >= 5
        ws.append([date, 100 if not weekend else 10, "휴일" if weekend else "평일"])
    return _save_wb(wb, "S9_주말포함.xlsx", SOURCE_DIR)


def source_10_dates_in_columns() -> str:
    """날짜가 '열' 방향으로 늘어선 표 (전치 필요)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "가로형"
    header = ["항목"] + [_dt.date(2026, 4, d).isoformat() for d in range(1, 11)]
    ws.append(header)
    ws.append(["사용량"] + [300 + d for d in range(1, 11)])
    ws.append(["온도"] + [16.0 + d * 0.1 for d in range(1, 11)])
    ws.append(["최대수요"] + [80 + d for d in range(1, 11)])
    return _save_wb(wb, "S10_열방향날짜.xlsx", SOURCE_DIR)


def source_11_many_rows() -> str:
    """템플릿이 예상한 것보다 행이 훨씬 많은 경우 (200행)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "대량"
    ws.append(["일자", "부서", "사용량"])
    base = _dt.date(2026, 5, 1)
    for index in range(200):
        ws.append([base + _dt.timedelta(days=index % 31), f"부서{index % 7}", 100 + index])
    return _save_wb(wb, "S11_행많음.xlsx", SOURCE_DIR)


def source_12_single_row() -> str:
    """데이터 행이 하나뿐인 경우."""
    wb = Workbook()
    ws = wb.active
    ws.title = "단일"
    ws.append(["부서", "담당자", "사용량", "최대수요", "온도", "비고"])
    ws.append(["총무팀", "정별", 990, 42, 21.5, "특이사항 없음"])
    return _save_wb(wb, "S12_한행.xlsx", SOURCE_DIR)


def source_13_weekday_suffix() -> str:
    """'2026-01-05(월)' 처럼 날짜 뒤에 요일이 덧붙은 표기."""
    wb = Workbook()
    ws = wb.active
    ws.title = "일별"
    ws.append(["일자", "사용량", "온도"])
    days = ["월", "화", "수", "목", "금", "토", "일"]
    for day in range(1, 11):
        date = _dt.date(2026, 7, day)
        weekday = days[date.weekday()]
        ws.append([f"2026-07-{day:02d}({weekday})", 200 + day, 19.0 + day * 0.1])
    return _save_wb(wb, "S13_날짜요일병기.xlsx", SOURCE_DIR)


def source_14_duplicate_dates() -> str:
    """같은 날짜에 측정치가 여러 번 기록된 경우 (오전/오후 검침 등)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "검침"
    ws.append(["일자", "구분", "사용량", "온도"])
    for day in range(1, 6):
        date = _dt.date(2026, 8, day)
        ws.append([date, "오전", 50 + day, 22.0 + day * 0.1])
        ws.append([date, "오후", 60 + day, 24.0 + day * 0.1])
    return _save_wb(wb, "S14_같은날짜중복.xlsx", SOURCE_DIR)


def source_15_units_and_negative() -> str:
    """단위 붙은 문자열 숫자와 음수(증감률 등)가 섞인 표."""
    wb = Workbook()
    ws = wb.active
    ws.title = "증감"
    ws.append(["일자", "사용량", "전일대비"])
    ws.append([_dt.date(2026, 9, 1), "1,250kWh", "-3.5%"])
    ws.append([_dt.date(2026, 9, 2), "1,180kWh", "-5.6%"])
    ws.append([_dt.date(2026, 9, 3), "1,320kWh", "11.9%"])
    ws.append([_dt.date(2026, 9, 4), -15, "-101.1%"])  # 설비 이상으로 음수 발생
    return _save_wb(wb, "S15_단위및음수.xlsx", SOURCE_DIR)


def source_16_row_and_col_offset() -> str:
    """제목 블록이 위쪽 여러 행을 차지하고, 표 자체도 왼쪽 여러 열을 비운 경우."""
    wb = Workbook()
    ws = wb.active
    ws.title = "보고서"
    ws["B2"] = "2026년 4분기 에너지 보고"
    ws["B2"].font = Font(size=14, bold=True)
    ws["B3"] = "작성부서: 시설관리팀"
    ws["D6"] = "부서"
    ws["E6"] = "사용량"
    ws["F6"] = "최대수요"
    rows = [("시설관리팀", 12500, 480), ("생산1팀", 30800, 910), ("생산2팀", 21450, 760)]
    for offset, record in enumerate(rows, start=7):
        for c_offset, value in enumerate(record):
            ws.cell(row=offset, column=4 + c_offset, value=value)
    return _save_wb(wb, "S16_행열오프셋.xlsx", SOURCE_DIR)


def source_17_merged_in_body() -> str:
    """데이터 본문 안에 병합 셀이 있는 경우 (부서명을 여러 날짜에 걸쳐 병합)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "부서별"
    ws.append(["부서", "일자", "사용량"])
    ws.merge_cells("A2:A4")
    ws["A2"] = "시설관리팀"
    for offset, day in enumerate((1, 2, 3), start=2):
        ws.cell(row=offset, column=2, value=_dt.date(2026, 10, day))
        ws.cell(row=offset, column=3, value=400 + day)
    ws.merge_cells("A5:A6")
    ws["A5"] = "생산1팀"
    for offset, day in enumerate((1, 2), start=5):
        ws.cell(row=offset, column=2, value=_dt.date(2026, 10, day))
        ws.cell(row=offset, column=3, value=900 + day)
    return _save_wb(wb, "S17_본문병합.xlsx", SOURCE_DIR)


def source_18_uncalculated_formula() -> str:
    """수식만 있고 '계산된 값' 캐시가 없는 원본 (프로그램이 만든 파일 등).

    openpyxl 로 값을 채워 넣기만 하고 한 번도 엑셀/LibreOffice 로 열어
    재계산하지 않은 상태를 그대로 재현한다. 여기 쓰인 ``=B2+C2`` 는 내장
    수식 계산기가 직접 계산할 수 있는 범위라, 이제는 오류 없이 정상 값이
    나와야 한다 (FormulaCacheError 는 계산기가 못 푸는 수식에서만 난다).
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "수식원본"
    ws.append(["일자", "기본", "부가", "합계"])
    ws.append([_dt.date(2026, 11, 1), 100, 20, "=B2+C2"])
    ws.append([_dt.date(2026, 11, 2), 110, 22, "=B3+C3"])
    return _save_wb(wb, "S18_계산안된수식.xlsx", SOURCE_DIR)


def source_19_full_month_no_gaps() -> str:
    """결측 없이 31일이 전부 채워진 달 (정상 케이스 대조군)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "일별"
    ws.append(["일자", "사용량", "최대수요", "온도"])
    for day in range(1, 32):
        ws.append([_dt.date(2026, 12, day), 300 + day, 90 + (day % 5), 10.0 + day * 0.05])
    return _save_wb(wb, "S19_전체31일.xlsx", SOURCE_DIR)


def source_20_totals_row_mixed() -> str:
    """일단위 데이터 맨 아래에 '합계' 같은 요약 행이 섞여 있는 경우.

    날짜로 읽히지 않는 행은 집계에서 자동 제외되어야 한다(스킵 카운트로 확인).
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "일별"
    ws.append(["일자", "사용량", "최대수요"])
    for day in range(1, 6):
        ws.append([_dt.date(2026, 5, day), 300 + day, 90 + day])
    ws.append(["합계", 1515, ""])  # 요약 행: 날짜가 아니므로 집계에서 스킵되어야 함
    ws.append(["작성자: 김하늘", "", ""])
    return _save_wb(wb, "S20_합계행혼재.xlsx", SOURCE_DIR)


def source_21_multi_block_sheet() -> str:
    """한 시트 안에 표가 여러 개(좌우로 3개 + 위아래로 별개 표) 뒤섞인 경우.

    대형 빌딩 에너지 관리 대장에서 흔한 구조: 날짜가 가로로 늘어선 표,
    전력/스팀/용수가 나란히 놓인 요약표, 그 아래 또 다른 일별 표가 이어짐.
    ``list_table_blocks()`` 가 이 덩어리들을 정확히 나누는지 검증하는 원본.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "에너지사용량"

    ws["A1"] = "일별값(kwh)"
    ws["C2"] = "2026-07-01"
    ws["A3"] = "일자"
    days = ["1일 수", "2일 목", "3일 금", "4일 토", "5일 일", "6일 월", "7일 화"]
    for index, day in enumerate(days):
        ws.cell(row=3, column=2 + index, value=day)
    ws["A4"] = "보고서값(2차)"
    ws["A5"] = "보고서값(1차)"
    # 6행: 빈 행 (표 구분)

    ws["B7"] = "전력"
    ws["B8"], ws["C8"], ws["D8"] = "구분", "사용량", "사용금액"
    ws["B9"], ws["C9"], ws["D9"] = "금년월보사용량", 5312566, 912061331

    ws["F7"] = "스팀"
    ws["F8"], ws["G8"], ws["H8"] = "구분", "사용량", "사용금액"
    ws["F9"], ws["G9"] = "금년월보사용량", 262000

    ws["J7"] = "용수"
    ws["J8"], ws["K8"] = "구분", "사용량"
    ws["J9"], ws["K9"] = "금년월보사용량", 29918
    # 13~15행: 빈 행 (표 구분)

    ws["A16"], ws["B16"] = "일별값", "Nm3"
    ws["A17"] = "일자"
    for index, day in enumerate(days):
        ws.cell(row=17, column=2 + index, value=day)
    ws["A18"] = "주방스팀(Nm3)"
    for index, value in enumerate([686, 673, 642, 181, 168, 651, 660]):
        ws.cell(row=18, column=2 + index, value=value)

    return _save_wb(wb, "S21_다중표시트.xlsx", SOURCE_DIR)


def source_22_uncalculated_header_formula() -> str:
    """날짜 헤더가 '하루씩 더하는 수식'인데 계산된 값이 없는 경우.

    대형빌딩 에너지 대장에서 실제로 나타난 패턴: 맨 왼쪽 날짜 하나만 값이고
    나머지는 ``=B3+1`` 식으로 하루씩 더하는 수식에 'd일 요일' 표시형식을
    입힌 것. 내장 수식 계산기가 이런 연쇄 덧셈은 직접 계산하므로, 이제는
    헤더 행이 정상적인 날짜로 채워져야 한다.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "일지"
    ws["B3"] = _dt.date(2026, 7, 1)
    ws["B3"].number_format = 'd"일" aaa'
    for col in range(3, 9):  # C~H: 하루씩 더하는 수식 (계산 캐시 없음)
        letter = chr(ord("A") + col - 1)
        prev_letter = chr(ord("A") + col - 2)
        cell = ws.cell(row=3, column=col, value=f"={prev_letter}3+1")
        cell.number_format = 'd"일" aaa'
    ws["A2"] = "일자"
    return _save_wb(wb, "S22_헤더수식미계산.xlsx", SOURCE_DIR)


def source_23_cross_sheet_formula() -> str:
    """다른 시트를 참조하는 수식에 계산 캐시가 없는 경우.

    ``formulas`` 패키지 기반 계산기는 워크북 전체를 계산 그래프로 만들기
    때문에, 시트 간 참조도 정상적으로 계산되어야 한다.
    """
    wb = Workbook()
    summary = wb.active
    summary.title = "요약"
    summary["A1"] = 999

    ws = wb.create_sheet("본표")
    ws.append(["일자", "기본", "요약값"])
    ws.append([_dt.date(2026, 12, 1), 100, "=요약!A1"])
    ws.append([_dt.date(2026, 12, 2), 110, "=요약!A1"])
    return _save_wb(wb, "S23_시트간참조.xlsx", SOURCE_DIR)


def source_24_truly_unsupported_formula() -> str:
    """계산기도 계산 못 하는 수식(존재하지 않는 외부 파일 참조)에 캐시도 없음.

    계산기가 아무리 좋아져도 못 푸는 경우는 항상 있다 - 디스크에 없는
    다른 엑셀 파일을 참조하는 수식이 대표적이다. 이런 경우 여전히
    FormulaCacheError 로 안전하게 걸러야 한다(틀린 값을 만들어내지 않고).
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "본표"
    ws.append(["일자", "기본", "외부참조"])
    ws.append([_dt.date(2026, 12, 1), 100, "='[없는파일.xlsx]Sheet1'!A1"])
    ws.append([_dt.date(2026, 12, 2), 110, "='[없는파일.xlsx]Sheet1'!A1"])
    return _save_wb(wb, "S24_미지원수식.xlsx", SOURCE_DIR)


# =========================================================================== #
# 실행
# =========================================================================== #
def _save_doc(doc: Document, name: str) -> str:
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    path = os.path.join(TEMPLATE_DIR, name)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = run.font.size or Pt(10)
    doc.save(path)
    return path


def _save_wb(wb: Workbook, name: str, directory: str = TEMPLATE_DIR) -> str:
    os.makedirs(directory, exist_ok=True)
    path = os.path.join(directory, name)
    wb.save(path)
    return path


TEMPLATE_BUILDERS = [
    word_1_plain,
    word_2_simple_table,
    word_3_merged_table,
    word_4_multi_table,
    word_5_repeated_tags,
    word_6_header_footer,
    word_7_loop_table,
    word_8_hyphen_tag,
    word_9_textbox,
    word_10_vertical_group_table,
    word_11_bulleted_list,
    word_12_split_runs,
    excel_1_single_sheet,
    excel_2_multi_sheet,
    excel_3_irregular,
    excel_4_styled,
    excel_5_formula,
    excel_6_table_anchor,
    excel_7_cell_coords,
    excel_8_frozen_and_rules,
    excel_9_seal_grid,
]

SOURCE_BUILDERS = [
    source_1_single_header,
    source_2_two_row_header,
    source_3_blanks,
    source_4_mixed_types,
    source_5_daily_one_month,
    source_6_daily_multi_month,
    source_7_date_formats,
    source_8_day_only,
    source_9_weekend,
    source_10_dates_in_columns,
    source_11_many_rows,
    source_12_single_row,
    source_13_weekday_suffix,
    source_14_duplicate_dates,
    source_15_units_and_negative,
    source_16_row_and_col_offset,
    source_17_merged_in_body,
    source_18_uncalculated_formula,
    source_19_full_month_no_gaps,
    source_20_totals_row_mixed,
    source_21_multi_block_sheet,
    source_22_uncalculated_header_formula,
    source_23_cross_sheet_formula,
    source_24_truly_unsupported_formula,
]


def build_all() -> dict[str, list[str]]:
    templates = [builder() for builder in TEMPLATE_BUILDERS]
    sources = [builder() for builder in SOURCE_BUILDERS]
    return {"templates": templates, "sources": sources}


if __name__ == "__main__":
    made = build_all()
    for group, paths in made.items():
        print(f"[{group}] {len(paths)}개")
        for path in paths:
            print("   ", os.path.relpath(path, HERE))
