"""워드(.docx) 템플릿 처리 - docxtpl 기반.

왜 그냥 docxtpl 을 쓰지 않는가
------------------------------
docxtpl 은 ``{{태그}}`` 안의 내용을 jinja2 식으로 해석한다. 그런데 실무에서
쓰는 태그 이름에는 ``{{연-월}}``, ``{{1월 실적}}`` 처럼 jinja2 가 연산자로
읽어버리는 글자가 흔히 들어간다. 그래서 여기서는 ``patch_xml`` 단계에 끼어들어

    {{연-월}}  ->  {{t_9f2c1a...}}

처럼 안전한 이름으로 바꾼 뒤, 같은 이름으로 값을 넘긴다. 사용자는 원래 이름
그대로 매핑 화면에서 보게 된다.

``{% for %}`` 같은 진짜 jinja 문법과 ``{{ r.사용량 }}`` 처럼 점/필터가 들어간
식은 건드리지 않으므로, docxtpl 을 이미 아는 사람은 그대로 쓸 수 있다.
"""

from __future__ import annotations

import hashlib
import os
import re
from typing import Any, Optional

from docx import Document
from docx.oxml import parse_xml
from docxtpl import DocxTemplate

from ..errors import TemplateError
from ..mapping import TemplateSlot
from .base import TAG_RE, TemplateHandler, is_simple_tag

__all__ = ["WordTemplate", "AliasingDocxTemplate"]

_ALIAS_PREFIX = "t_"


def _alias_for(expression: str) -> str:
    """태그 이름 -> jinja2 가 받아들이는 안전한 변수명 (항상 같은 값)."""
    digest = hashlib.md5(expression.strip().encode("utf-8")).hexdigest()[:16]
    return f"{_ALIAS_PREFIX}{digest}"


class AliasingDocxTemplate(DocxTemplate):
    """``patch_xml`` 에 태그 치환을 끼워 넣은 DocxTemplate.

    ``patch_xml`` 은 본문/머리글/바닥글/각주에 모두 공통으로 불리므로, 여기 한
    군데만 손보면 문서 전체에 일관되게 적용된다.
    """

    def patch_xml(self, src_xml: str) -> str:
        xml = super().patch_xml(src_xml)
        return _rewrite_simple_tags(xml)


def _rewrite_simple_tags(xml: str) -> str:
    def replace(match: "re.Match[str]") -> str:
        inner = match.group(1)
        if not is_simple_tag(inner):
            return match.group(0)
        return "{{ %s }}" % _alias_for(inner)

    return TAG_RE.sub(replace, xml)


class WordTemplate(TemplateHandler):
    kind = "word"
    extension = ".docx"

    def __init__(self, path: str) -> None:
        super().__init__(path)
        self._slots: Optional[list[TemplateSlot]] = None

    # ------------------------------------------------------------------ #
    # 스캔
    # ------------------------------------------------------------------ #
    def scan(self) -> list[TemplateSlot]:
        if self._slots is not None:
            return self._slots

        try:
            document = Document(self.path)
        except Exception as exc:  # noqa: BLE001
            raise TemplateError(
                f"워드 템플릿을 열지 못했습니다: {self.name}",
                f"파일이 손상되었거나 .docx 가 아닐 수 있습니다. ({exc})",
            ) from exc

        found: dict[str, TemplateSlot] = {}

        for where, text in _iter_text_locations(document):
            for match in TAG_RE.finditer(text):
                inner = match.group(1).strip()
                if not is_simple_tag(inner):
                    continue
                slot = found.get(inner)
                if slot is None:
                    found[inner] = TemplateSlot(
                        key=inner,
                        kind="tag",
                        where=where,
                        sample=_snippet(text, match.start(), match.end()),
                        occurrences=1,
                    )
                else:
                    slot.occurrences += 1
                    if where not in slot.where:
                        slot.where = f"{slot.where}, {where}"

        # python-docx 의 문단/표 순회는 텍스트 상자·도형 안을 보지 못한다.
        # 렌더링은 XML 전체를 대상으로 하므로, 스캔도 XML 을 한 번 훑어서
        # 빠진 태그가 없는지 확인한다. (놓치면 조용히 빈칸으로 나가 버린다.)
        for key in _scan_raw_xml(self.path):
            if key not in found:
                found[key] = TemplateSlot(
                    key=key,
                    kind="tag",
                    where="본문 외 영역(텍스트 상자·도형 등)",
                    sample="",
                    occurrences=1,
                )

        self._slots = list(found.values())
        return self._slots

    # ------------------------------------------------------------------ #
    # 렌더링
    # ------------------------------------------------------------------ #
    def render(
        self,
        context: dict[str, Any],
        output_path: str,
        table_data: Optional[dict] = None,
    ) -> str:
        try:
            template = AliasingDocxTemplate(self.path)
        except Exception as exc:  # noqa: BLE001
            raise TemplateError(
                f"워드 템플릿을 열지 못했습니다: {self.name}",
                f"다른 프로그램이 파일을 열고 있지 않은지 확인해 주세요. ({exc})",
            ) from exc

        render_context = _build_render_context(context, table_data)

        try:
            template.render(render_context)
        except Exception as exc:  # noqa: BLE001
            raise TemplateError(
                f"워드 템플릿을 채우는 중 오류가 발생했습니다: {self.name}",
                _render_hint(exc),
            ) from exc

        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        try:
            template.save(output_path)
        except PermissionError as exc:
            raise TemplateError(
                f"결과 파일을 저장하지 못했습니다: {os.path.basename(output_path)}",
                "같은 이름의 파일이 워드에서 열려 있으면 닫은 뒤 다시 시도해 주세요.",
            ) from exc
        return os.path.abspath(output_path)


# --------------------------------------------------------------------------- #
# 보조
# --------------------------------------------------------------------------- #
def _build_render_context(
    context: dict[str, Any], table_data: Optional[dict]
) -> dict[str, Any]:
    """사용자 태그 -> 별칭으로 바꾸고, 반복문용 변수도 함께 넣는다."""
    render_context: dict[str, Any] = {}
    for key, value in context.items():
        render_context[_alias_for(key)] = value
        # 원래 이름이 그대로 유효한 변수명이면 그 이름으로도 접근할 수 있게 둔다.
        if key.isidentifier():
            render_context.setdefault(key, value)

    rows = (table_data or {}).get("rows") or []
    columns = (table_data or {}).get("columns") or []
    render_context.setdefault("rows", rows)
    render_context.setdefault("표", rows)
    render_context.setdefault("columns", columns)
    render_context.setdefault("컬럼", columns)
    return render_context


def _scan_raw_xml(path: str) -> set[str]:
    """문서 XML(본문 + 머리글/바닥글) 전체에서 단순 태그 이름을 모은다.

    docxtpl 의 ``patch_xml`` 을 먼저 돌려서, 워드가 태그를 여러 run 으로
    쪼개 놓은 경우에도 온전한 이름으로 읽히게 한다.
    """
    keys: set[str] = set()
    try:
        template = DocxTemplate(path)
        template.init_docx()
        parts = [template.patch_xml(template.get_xml())]
        for uri in (DocxTemplate.HEADER_URI, DocxTemplate.FOOTER_URI):
            for _rel_key, part in template.get_headers_footers(uri):
                parts.append(
                    template.patch_xml(template.xml_to_string(parse_xml(part.blob)))
                )
    except Exception:  # noqa: BLE001 - 보조 스캔이므로 실패해도 본 스캔 결과를 쓴다
        return keys

    for xml in parts:
        for match in TAG_RE.finditer(xml):
            inner = match.group(1).strip()
            if is_simple_tag(inner):
                keys.add(inner)
    return keys


def _iter_text_locations(document) -> list[tuple[str, str]]:
    """(위치 설명, 텍스트) 목록. 본문·표·머리글·바닥글을 모두 훑는다.

    문단 단위가 아니라 '문단 전체 텍스트'로 훑기 때문에, 워드가 태그를 여러
    run 으로 쪼개 놓아도 스캔에서 놓치지 않는다.
    """
    out: list[tuple[str, str]] = []

    for index, paragraph in enumerate(document.paragraphs, start=1):
        if "{{" in paragraph.text:
            out.append((f"본문 {index}번째 문단", paragraph.text))

    for t_index, table in enumerate(document.tables, start=1):
        out.extend(_iter_table(table, f"표{t_index}"))

    for s_index, section in enumerate(document.sections, start=1):
        for label, part in (
            ("머리글", section.header),
            ("바닥글", section.footer),
            ("첫 페이지 머리글", section.first_page_header),
            ("첫 페이지 바닥글", section.first_page_footer),
            ("짝수 페이지 머리글", section.even_page_header),
            ("짝수 페이지 바닥글", section.even_page_footer),
        ):
            if part is None:
                continue
            prefix = label if len(document.sections) == 1 else f"{s_index}구역 {label}"
            for paragraph in part.paragraphs:
                if "{{" in paragraph.text:
                    out.append((prefix, paragraph.text))
            for t_index, table in enumerate(part.tables, start=1):
                out.extend(_iter_table(table, f"{prefix} 표{t_index}"))

    return out


def _iter_table(table, label: str) -> list[tuple[str, str]]:
    """표 안의 셀을 훑는다. 중첩 표와 병합 셀도 포함.

    병합된 셀은 python-docx 가 같은 셀 객체를 여러 좌표에서 돌려주므로,
    이미 본 셀은 건너뛰어 중복 카운트를 막는다.

    주의: lxml 의 엘리먼트 프록시는 참조가 없어지면 회수되고 ``id()`` 값이
    재사용될 수 있다. 그래서 본 엘리먼트를 ``anchors`` 에 붙잡아 둔다.
    """
    out: list[tuple[str, str]] = []
    seen: set[int] = set()
    anchors: list[Any] = []  # GC 로 id() 가 재사용되지 않도록 참조를 유지
    for r_index, row in enumerate(table.rows, start=1):
        for c_index, cell in enumerate(row.cells, start=1):
            element = cell._tc
            marker = id(element)
            if marker in seen:
                continue
            seen.add(marker)
            anchors.append(element)
            where = f"{label} {r_index}행{c_index}열"
            for paragraph in cell.paragraphs:
                if "{{" in paragraph.text:
                    out.append((where, paragraph.text))
            for n_index, nested in enumerate(cell.tables, start=1):
                out.extend(_iter_table(nested, f"{where} 중첩표{n_index}"))
    return out


def _snippet(text: str, start: int, end: int, width: int = 24) -> str:
    left = max(0, start - width)
    right = min(len(text), end + width)
    prefix = "…" if left > 0 else ""
    suffix = "…" if right < len(text) else ""
    return f"{prefix}{text[left:right].strip()}{suffix}"


def _render_hint(exc: Exception) -> str:
    message = str(exc)
    if "unexpected" in message.lower() or "syntax" in message.lower():
        return (
            "템플릿 안에 {% ... %} 문법이 잘못 적혀 있을 수 있습니다. "
            f"원문: {message}"
        )
    return f"원인: {message}"
