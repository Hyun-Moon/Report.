#!/usr/bin/env python3
"""``templates/`` 폴더에 바로 쓸 수 있는 기본 템플릿을 만든다.

    python tools/make_starter_templates.py

이미 같은 이름의 파일이 있으면 건드리지 않는다(사용자가 고친 것을 덮지 않기 위함).
검증용 샘플과 달리, 이쪽은 '실제로 배포되는 시작용 양식'이다.
"""

from __future__ import annotations

import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATES = os.path.join(ROOT, "templates")

THIN = Side(style="thin", color="9E9E9E")
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
HEAD = PatternFill("solid", fgColor="E8EEF7")
CENTER = Alignment(horizontal="center", vertical="center")


def word_monthly_report(path: str) -> None:
    doc = Document()
    title = doc.add_heading("월간 실적 보고서", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section = doc.sections[0]
    section.header.paragraphs[0].text = "{{부서}}"
    section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    section.footer.paragraphs[0].text = "작성 {{담당자}} · 출력 {{오늘}}"
    section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_table(rows=2, cols=4)
    info.style = "Table Grid"
    info.cell(0, 0).text = "대상 기간"
    info.cell(0, 1).text = "{{대상월}}"
    info.cell(0, 2).text = "작성일"
    info.cell(0, 3).text = "{{오늘}}"
    info.cell(1, 0).text = "작성자"
    info.cell(1, 1).text = "{{담당자}}"
    info.cell(1, 2).text = "집계 일수"
    info.cell(1, 3).text = "{{집계일수}}"

    doc.add_paragraph("")
    doc.add_paragraph("1. 실적 요약")
    summary = doc.add_table(rows=4, cols=3)
    summary.style = "Table Grid"
    for row, values in zip(
        summary.rows,
        [
            ("항목", "값", "단위"),
            ("사용량", "{{사용량}}", "kWh"),
            ("최대수요", "{{최대수요}}", "kW"),
            ("평균온도", "{{온도}}", "℃"),
        ],
    ):
        for cell, value in zip(row.cells, values):
            cell.text = value

    doc.add_paragraph("")
    doc.add_paragraph("2. 월별 상세")
    detail = doc.add_table(rows=4, cols=3)
    detail.style = "Table Grid"
    detail.cell(0, 0).text = "연-월"
    detail.cell(0, 1).text = "일수"
    detail.cell(0, 2).text = "사용량"
    detail.cell(1, 0).text = "{%tr for r in rows %}"
    detail.cell(2, 0).text = "{{ r[0] }}"
    detail.cell(2, 1).text = "{{ r[1] }}"
    detail.cell(2, 2).text = "{{ r[2] }}"
    detail.cell(3, 0).text = "{%tr endfor %}"

    doc.add_paragraph("")
    doc.add_paragraph("3. 특이사항")
    note = doc.add_table(rows=1, cols=1)
    note.style = "Table Grid"
    note.cell(0, 0).text = "{{비고}}"

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = run.font.size or Pt(10)
    doc.save(path)


def word_simple_memo(path: str) -> None:
    doc = Document()
    doc.add_heading("업무 보고", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("수신: {{부서}}")
    doc.add_paragraph("작성: {{담당자}}   작성일: {{오늘}}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "{{대상월}} 실적을 아래와 같이 보고합니다. 총 사용량은 {{사용량}} kWh, "
        "최대수요는 {{최대수요}} kW, 평균온도는 {{온도}} ℃ 입니다."
    )
    doc.add_paragraph("")
    doc.add_paragraph("특이사항: {{비고}}")
    doc.save(path)


def excel_summary(path: str) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "요약"

    ws.merge_cells("A1:D1")
    ws["A1"] = "{{대상월}} 실적 요약"
    ws["A1"].font = Font(size=14, bold=True)
    ws["A1"].alignment = CENTER
    ws["A1"].fill = HEAD

    ws["A3"], ws["B3"] = "작성자", "{{담당자}}"
    ws["C3"], ws["D3"] = "집계 일수", "{{집계일수}}"

    for index, name in enumerate(["항목", "값", "단위", "비고"], start=1):
        cell = ws.cell(row=5, column=index, value=name)
        cell.font = Font(bold=True)
        cell.fill = HEAD
        cell.border = BOX
        cell.alignment = CENTER

    rows = [
        ("사용량", "{{사용량}}", "kWh", ""),
        ("최대수요", "{{최대수요}}", "kW", ""),
        ("평균온도", "{{온도}}", "℃", "{{비고}}"),
    ]
    for r_offset, record in enumerate(rows, start=6):
        for c_offset, value in enumerate(record, start=1):
            cell = ws.cell(row=r_offset, column=c_offset, value=value)
            cell.border = BOX
            if c_offset == 2:
                cell.number_format = "#,##0.##"

    ws["A10"] = "합계(수식)"
    ws["B10"] = "=SUM(B6:B7)"
    ws["A10"].font = Font(bold=True)

    for column, width in zip("ABCD", (16, 18, 10, 24)):
        ws.column_dimensions[column].width = width

    detail = wb.create_sheet("월별상세")
    detail["A1"] = "{{대상월}} 월별 상세"
    detail["A1"].font = Font(bold=True)
    detail["A3"] = "{{#표}}"
    for column in range(1, 8):
        detail.column_dimensions[get_column_letter(column)].width = 14

    wb.save(path)


BUILDERS = {
    "기본_월간실적보고서.docx": word_monthly_report,
    "기본_업무보고_간단.docx": word_simple_memo,
    "기본_실적요약표.xlsx": excel_summary,
}


def main() -> int:
    os.makedirs(TEMPLATES, exist_ok=True)
    made, skipped = [], []
    for name, builder in BUILDERS.items():
        path = os.path.join(TEMPLATES, name)
        if os.path.exists(path):
            skipped.append(name)
            continue
        builder(path)
        made.append(name)

    for name in made:
        print(f"만듦   templates/{name}")
    for name in skipped:
        print(f"건너뜀 templates/{name} (이미 있음)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
